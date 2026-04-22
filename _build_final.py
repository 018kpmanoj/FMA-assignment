"""
FMA EC1 Assignment - Final Build Script
HCL Technologies vs Wipro | Data: HCL Tech-1.xlsx, Wipro-1.xlsx (Moneycontrol)
Students: K P Manoj (2025MB26539), T Venkata Suresh (2025MB26590), Japleen Kaur Khorana (2025MB26602), Varsha PS (2025MB26503), Prashanth G (2025MB26598)
"""
import os, shutil, copy
from pathlib import Path
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter
from openpyxl.workbook.properties import CalcProperties
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = Path(__file__).resolve().parent
FINAL = BASE

YEARS = ["FY2025", "FY2024", "FY2023", "FY2022", "FY2021"]
FACE_VALUE = 2

# ========================  RAW DATA (Consolidated)  ========================
# Source: HCL Tech-1.xlsx  (Moneycontrol, consolidated columns C10-C14)
H_PNL = dict(
    revenue    = [117055, 109913, 101456, 85651, 75379],
    other_inc  = [2485, 1495, 1358, 1067, 927],
    total_rev  = [119540, 111408, 102814, 86718, 76306],
    purchase   = [1976, 1754, 2072, 1473, 1698],
    op_exp     = [15162, 14578, 14950, 12515, 10158],
    inv_change = [52, 43, -67, -67, -3],
    emp_exp    = [66755, 62480, 55280, 46130, 38853],
    fin_cost   = [644, 553, 353, 319, 511],
    depr       = [4084, 4173, 4145, 4326, 4611],
    other_exp  = [7606, 6860, 6593, 5070, 4625],
    total_exp  = [96279, 90441, 83326, 69766, 60453],
    pbt        = [23261, 20967, 19488, 16952, 15853],
    tax        = [5862, 5257, 4643, 3428, 4684],
    pat        = [17399, 15710, 14845, 13524, 11169],
    pat_mi     = [17390, 15702, 14851, 13499, 11145],
)
H_BS = dict(
    eq_capital   = [543, 543, 543, 543, 543],
    reserves     = [69112, 67720, 64862, 61371, 59370],
    total_shf    = [69655, 68263, 65405, 61914, 59913],
    lt_borr      = [70, 2223, 2111, 3923, 3828],
    st_borr      = [2221, 104, 140, 62, 0],
    trade_pay    = [13234, 5853, 6428, 6278, 1726],
    total_cl     = [28039, 22726, 21431, 18775, 17383],
    total_ncl    = [7832, 8780, 6582, 8252, 8729],
    fixed_assets = [14475, 15039, 16092, 17789, 20265],
    inventories  = [133, 185, 228, 161, 94],
    trade_rec    = [25842, 25521, 25506, 20671, 13663],
    cash         = [21289, 20150, 14724, 12636, 8888],
    curr_inv     = [7473, 7043, 5385, 6239, 6773],
    total_ca     = [62109, 59331, 53577, 48041, 43051],
    total_assets = [105544, 99777, 93411, 89033, 86194],
)
H_MKT = dict(price_bv=[6.20, 6.14, 4.51, 5.10, 4.46],
              bv_share=[256.56, 251.46, 240.88, 228.38, 221.30])

# Source: Wipro-1.xlsx  (Moneycontrol, consolidated columns C10-C14)
W_PNL = dict(
    revenue    = [89088.4, 89760.3, 90487.6, 79093.4, 61943],
    other_inc  = [3884, 2630.8, 2265.7, 2061.2, 2390.7],
    total_rev  = [92972.4, 92391.1, 92753.3, 81373.2, 64325.6],
    purchase   = [296.7, 383.2, 649.4, 673.5, 695.7],
    op_exp     = [12332.8, 14084.2, 15336.7, 15449.9, 12075.2],
    inv_change = [19.5, 27.8, 15, -36.9, 31.5],
    emp_exp    = [53347.7, 54930.1, 53764.4, 45007.5, 33237.1],
    fin_cost   = [1477, 1255.2, 1007.7, 532.5, 508.8],
    depr       = [2957.9, 3407.1, 3340.2, 3077.8, 2763.4],
    other_exp  = [5070.5, 3559.2, 3868.5, 1533.8, 1124],
    total_exp  = [75502.1, 77646.8, 77981.9, 66238.1, 50435.7],
    pbt        = [17470.3, 14744.3, 14771.4, 15135.1, 13889.9],
    tax        = [4277.7, 3608.9, 3399.2, 2897.4, 3034.9],
    pat        = [13192.6, 11135.4, 11372.2, 12237.7, 10855],
    pat_mi     = [13135.4, 11045.2, 11350, 12229.6, 10796.4],
)
W_BS = dict(
    eq_capital   = [2094.4, 1045, 1097.6, 1096.4, 1095.8],
    reserves     = [80269.7, 73488, 76570.3, 64306.6, 53805.2],
    total_shf    = [82364.1, 74533, 77667.9, 65403, 54901],
    lt_borr      = [6395.4, 6230, 6127.2, 5646.3, 745.8],
    st_borr      = [9786.3, 7916.6, 8882.1, 9523.3, 6036.3],
    trade_pay    = [5866.7, 5765.5, 5972.3, 6252.2, 5417.4],
    total_cl     = [28625.3, 25245.8, 26775.3, 30832.9, 23004],
    total_ncl    = [16982, 14877.8, 12631.6, 11218, 4677.4],
    fixed_assets = [13348.5, 13206.5, 15025.4, 15305, 11378.8],
    inventories  = [69.4, 90.7, 118.8, 133.4, 106.4],
    trade_rec    = [11774.5, 17382.2, 18686.5, 11521.9, 9429.8],
    cash         = [12197.4, 9695.3, 9188, 10383.6, 16979.3],
    curr_inv     = [41147.4, 31117.1, 30923.2, 24165.5, 17570.7],
    total_ca     = [77777.5, 65066.2, 66109.6, 62075.2, 52318.6],
    total_assets = [128185.2, 114790.6, 117133.7, 107505.4, 82732.2],
)
W_MKT = dict(price_bv=[3.33, 3.37, 2.58, 4.96, 4.13],
              bv_share=[78.65, 142.65, 141.63, 119.40, 100.48])

# FY2020 data for proper CCC averaging at FY2021 boundary
H_FY2020 = dict(inventories=91, trade_rec=14131, trade_pay=1166)
W_FY2020 = dict(inventories=186.5, trade_rec=10447.4, trade_pay=5840)

# =======================  RATIO COMPUTATION (Python)  =======================
def _avg(lst, i):
    return (lst[i] + lst[i+1]) / 2 if i < len(lst) - 1 else lst[i]

def compute_ratios(pnl, bs, mkt):
    n = 5
    rev = pnl['revenue']
    td  = [bs['lt_borr'][i] + bs['st_borr'][i] for i in range(n)]
    ce  = [bs['total_assets'][i] - bs['total_cl'][i] for i in range(n)]
    sh  = [bs['eq_capital'][i] / FACE_VALUE for i in range(n)]
    mp  = [mkt['price_bv'][i] * mkt['bv_share'][i] for i in range(n)]
    cogs = [pnl['purchase'][i] + pnl['op_exp'][i] + pnl['inv_change'][i] for i in range(n)]
    purch = [pnl['purchase'][i] + pnl['op_exp'][i] + pnl['other_exp'][i] for i in range(n)]
    ebit = [pnl['pbt'][i] + pnl['fin_cost'][i] for i in range(n)]
    r = {}
    r['Current Ratio']            = [bs['total_ca'][i]/bs['total_cl'][i] for i in range(n)]
    r['Quick Ratio']              = [(bs['total_ca'][i]-bs['inventories'][i])/bs['total_cl'][i] for i in range(n)]
    r['Absolute Liquid Ratio']    = [(bs['cash'][i]+bs['curr_inv'][i])/bs['total_cl'][i] for i in range(n)]
    r['Gross Profit Ratio (%)']   = [(rev[i]-cogs[i])/rev[i]*100 for i in range(n)]
    r['Operating Profit Ratio (%)'] = [ebit[i]/rev[i]*100 for i in range(n)]
    r['Net Profit Ratio (%)']     = [pnl['pat'][i]/rev[i]*100 for i in range(n)]
    r['Return on Equity (%)']     = [pnl['pat'][i]/_avg(bs['total_shf'],i)*100 for i in range(n)]
    r['Return on Capital Employed (%)'] = [ebit[i]/ce[i]*100 for i in range(n)]
    r['Earnings per Share (Rs)']  = [pnl['pat_mi'][i]/sh[i] for i in range(n)]
    r['Debt-Equity Ratio']        = [td[i]/bs['total_shf'][i] for i in range(n)]
    r['Interest Coverage Ratio']  = [ebit[i]/pnl['fin_cost'][i] for i in range(n)]
    r['Debtors Turnover Ratio']   = [rev[i]/_avg(bs['trade_rec'],i) for i in range(n)]
    r['Avg Collection Period (days)'] = [365/r['Debtors Turnover Ratio'][i] for i in range(n)]
    r['Creditors Turnover Ratio'] = [purch[i]/_avg(bs['trade_pay'],i) for i in range(n)]
    r['Avg Payment Period (days)']= [365/r['Creditors Turnover Ratio'][i] for i in range(n)]
    r['Inventory Turnover Ratio'] = [cogs[i]/_avg(bs['inventories'],i) for i in range(n)]
    r['Inventory Holding Days']   = [365/r['Inventory Turnover Ratio'][i] for i in range(n)]
    r['Fixed Asset Turnover']     = [rev[i]/_avg(bs['fixed_assets'],i) for i in range(n)]
    r['P/E Ratio']                = [mp[i]/(pnl['pat_mi'][i]/sh[i]) for i in range(n)]
    r['_npm']  = [pnl['pat'][i]/rev[i] for i in range(n)]
    r['_at']   = [rev[i]/bs['total_assets'][i] for i in range(n)]
    r['_em']   = [bs['total_assets'][i]/bs['total_shf'][i] for i in range(n)]
    r['_roe_d']= [r['_npm'][i]*r['_at'][i]*r['_em'][i]*100 for i in range(n)]
    r['_dio']  = [365*_avg(bs['inventories'],i)/cogs[i] if cogs[i] else 0 for i in range(n)]
    r['_dso']  = [365*_avg(bs['trade_rec'],i)/rev[i] for i in range(n)]
    r['_dpo']  = [365*_avg(bs['trade_pay'],i)/purch[i] for i in range(n)]
    r['_ccc']  = [r['_dio'][i]+r['_dso'][i]-r['_dpo'][i] for i in range(n)]
    r['_mp']   = mp; r['_td']=td; r['_ce']=ce; r['_sh']=sh; r['_ebit']=ebit
    r['_cogs'] = cogs; r['_purch'] = purch
    return r

HR = compute_ratios(H_PNL, H_BS, H_MKT)
WR = compute_ratios(W_PNL, W_BS, W_MKT)

def _fix_ccc_fy2021(r, pnl, bs, fy2020):
    """Override FY2021 CCC components with proper 2-year averages using FY2020 data."""
    i = 4
    cogs = pnl['purchase'][i] + pnl['op_exp'][i] + pnl['inv_change'][i]
    purch = pnl['purchase'][i] + pnl['op_exp'][i] + pnl['other_exp'][i]
    rev = pnl['revenue'][i]
    avg_inv = (bs['inventories'][i] + fy2020['inventories']) / 2
    avg_rec = (bs['trade_rec'][i] + fy2020['trade_rec']) / 2
    avg_pay = (bs['trade_pay'][i] + fy2020['trade_pay']) / 2
    r['_dio'][i] = 365 * avg_inv / cogs if cogs else 0
    r['_dso'][i] = 365 * avg_rec / rev
    r['_dpo'][i] = 365 * avg_pay / purch
    r['_ccc'][i] = r['_dio'][i] + r['_dso'][i] - r['_dpo'][i]

_fix_ccc_fy2021(HR, H_PNL, H_BS, H_FY2020)
_fix_ccc_fy2021(WR, W_PNL, W_BS, W_FY2020)


# =======================  DATASET CLEANUP  =======================
def clean_dataset_files():
    """Remove standalone columns (1-8) and Ratios sheet from raw dataset files.

    The raw Moneycontrol files have:
      Columns 1-6 : Standalone financials
      Columns 7-8 : Empty separator
      Columns 9-14: Consolidated financials
    After cleanup, only consolidated data remains (columns 9-14 become 1-6).
    The Ratios sheet is removed because we compute our own ratios in excel_report.
    """
    for fname in ["HCL Tech-1.xlsx", "Wipro-1.xlsx"]:
        src = BASE / fname
        if not src.exists():
            print(f"[SKIP] {fname} not found")
            continue

        wb = load_workbook(src)

        if 'Ratios' in wb.sheetnames:
            del wb['Ratios']
            print(f"  Removed 'Ratios' sheet from {fname}")

        for ws in wb.worksheets:
            ws.delete_cols(1, 8)

        wb.save(src)
        wb.close()
        print(f"[OK] Cleaned {fname}: standalone columns removed, consolidated only")


# ==========================  EXCEL BUILDER  ==========================
THIN = Side(style='thin')
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
HDR_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
HDR_FONT = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
SEC_FILL = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
SEC_FONT = Font(name="Calibri", size=10, bold=True, color="1F4E79")
DATA_FONT = Font(name="Calibri", size=10)
NUM_FMT = '#,##0.00'
NUM_FMT_INT = '#,##0'

def _style_header(ws, row, max_col):
    for c in range(1, max_col+1):
        cell = ws.cell(row=row, column=c)
        cell.font = HDR_FONT; cell.fill = HDR_FILL
        cell.alignment = Alignment(horizontal='center', wrap_text=True)
        cell.border = BORDER

def _style_data(ws, row, max_col, fmt=NUM_FMT):
    for c in range(1, max_col+1):
        cell = ws.cell(row=row, column=c)
        cell.font = DATA_FONT; cell.border = BORDER
        if c > 1:
            cell.number_format = fmt
            cell.alignment = Alignment(horizontal='right')
        else:
            cell.alignment = Alignment(wrap_text=True)

def _style_section(ws, row, max_col):
    for c in range(1, max_col+1):
        cell = ws.cell(row=row, column=c)
        cell.font = SEC_FONT; cell.fill = SEC_FILL; cell.border = BORDER


def _write_full_data_sheet(ws, title, raw_file_path):
    """Copy ALL consolidated data from cleaned raw file into an excel_report sheet.

    After cleanup, the raw file has: Col1=Labels, Col2-6=FY2025..FY2021.
    Each sheet (Balance Sheet, P&L, Cash Flow) is copied in full.
    """
    ws.merge_cells('A1:F1')
    ws['A1'] = title
    ws['A1'].font = Font(name="Calibri", size=12, bold=True, color="1F4E79")

    raw_wb = load_workbook(raw_file_path, data_only=True)
    out_row = 3

    for sheet_name in raw_wb.sheetnames:
        raw_ws = raw_wb[sheet_name]

        ws.cell(out_row, 1, sheet_name.upper())
        _style_section(ws, out_row, 6)
        out_row += 1

        headers = ["Particulars"] + YEARS
        for c, h in enumerate(headers, 1):
            ws.cell(out_row, c, h)
        _style_header(ws, out_row, 6)
        out_row += 1

        for src_r in range(4, raw_ws.max_row + 1):
            label = raw_ws.cell(src_r, 1).value
            vals = [raw_ws.cell(src_r, c).value for c in range(2, 7)]

            if label is None and all(v is None for v in vals):
                continue

            ws.cell(out_row, 1, label)
            for ci, v in enumerate(vals, 2):
                if v is not None:
                    ws.cell(out_row, ci, v)

            if all(v is None for v in vals) and label:
                _style_section(ws, out_row, 6)
            else:
                _style_data(ws, out_row, 6)
            out_row += 1

        out_row += 1

    raw_wb.close()

    ws.column_dimensions['A'].width = 42
    for c in range(2, 7):
        ws.column_dimensions[get_column_letter(c)].width = 15


def _write_ratios_sheet(ws, hr, wr):
    ws.merge_cells('A1:L1')
    ws['A1'] = "Financial Ratios - HCL Technologies vs Wipro (FY2021-FY2025)"
    ws['A1'].font = Font(name="Calibri", size=12, bold=True, color="1F4E79")

    r = 3
    ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
    ws.cell(r, 2, "HCL Technologies").font = Font(bold=True, size=10)
    ws.cell(r, 2).alignment = Alignment(horizontal='center')
    ws.merge_cells(start_row=r, start_column=8, end_row=r, end_column=12)
    ws.cell(r, 8, "Wipro").font = Font(bold=True, size=10)
    ws.cell(r, 8).alignment = Alignment(horizontal='center')

    r = 4
    row4 = ["Ratio"] + YEARS + [""] + YEARS
    for c, h in enumerate(row4, 1):
        ws.cell(r, c, h)
    _style_header(ws, r, 12)

    RATIO_NAMES = [
        None,
        "Current Ratio", "Quick Ratio", "Absolute Liquid Ratio",
        None,
        "Gross Profit Ratio (%)", "Operating Profit Ratio (%)",
        "Net Profit Ratio (%)", "Return on Equity (%)",
        "Return on Capital Employed (%)", "Earnings per Share (Rs)",
        None,
        "Debt-Equity Ratio", "Interest Coverage Ratio",
        None,
        "Debtors Turnover Ratio", "Avg Collection Period (days)",
        "Creditors Turnover Ratio", "Avg Payment Period (days)",
        "Inventory Turnover Ratio", "Inventory Holding Days",
        "Fixed Asset Turnover",
        None,
        "P/E Ratio",
    ]
    SECTIONS = {
        0: "LIQUIDITY RATIOS",
        4: "PROFITABILITY RATIOS",
        11: "SOLVENCY RATIOS",
        14: "TURNOVER RATIOS",
        22: "VALUATION RATIOS",
    }

    r = 5
    for idx, name in enumerate(RATIO_NAMES):
        if idx in SECTIONS:
            ws.cell(r, 1, SECTIONS[idx])
            _style_section(ws, r, 12)
            r += 1
        if name is None:
            continue
        ws.cell(r, 1, name)
        for yi in range(5):
            ws.cell(r, 2+yi, round(hr[name][yi], 2))
            ws.cell(r, 8+yi, round(wr[name][yi], 2))
        _style_data(ws, r, 12)
        for c in range(8, 13):
            ws.cell(r, c).number_format = NUM_FMT
            ws.cell(r, c).border = BORDER
            ws.cell(r, c).alignment = Alignment(horizontal='right')
        r += 1

    ws.column_dimensions['A'].width = 34
    for c in range(2, 13):
        ws.column_dimensions[get_column_letter(c)].width = 13
    ws.column_dimensions['G'].width = 2


def _write_dupont_sheet(ws, hr, wr):
    ws.merge_cells('A1:L1')
    ws['A1'] = "DuPont Analysis - HCL Technologies vs Wipro"
    ws['A1'].font = Font(name="Calibri", size=12, bold=True, color="1F4E79")
    r = 3
    ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
    ws.cell(r, 2, "HCL Technologies").font = Font(bold=True)
    ws.cell(r, 2).alignment = Alignment(horizontal='center')
    ws.merge_cells(start_row=r, start_column=8, end_row=r, end_column=12)
    ws.cell(r, 8, "Wipro").font = Font(bold=True)
    ws.cell(r, 8).alignment = Alignment(horizontal='center')
    r = 4
    for c, h in enumerate(["Component"] + YEARS + [""] + YEARS, 1):
        ws.cell(r, c, h)
    _style_header(ws, r, 12)

    components = [
        ("Net Profit Margin", '_npm'),
        ("Asset Turnover", '_at'),
        ("Equity Multiplier", '_em'),
    ]
    r = 5
    for label, key in components:
        ws.cell(r, 1, label)
        for yi in range(5):
            ws.cell(r, 2+yi, round(hr[key][yi], 4))
            ws.cell(r, 8+yi, round(wr[key][yi], 4))
        _style_data(ws, r, 12, '0.0000')
        for c in range(8, 13):
            ws.cell(r, c).number_format = '0.0000'
            ws.cell(r, c).border = BORDER
        r += 1

    r += 1
    ws.cell(r, 1, "ROE via DuPont (%)")
    ws.cell(r, 1).font = Font(bold=True, size=10, color="1F4E79")
    for yi in range(5):
        ws.cell(r, 2+yi, round(hr['_roe_d'][yi], 2))
        ws.cell(r, 8+yi, round(wr['_roe_d'][yi], 2))
    _style_data(ws, r, 12)
    for c in range(8, 13):
        ws.cell(r, c).number_format = NUM_FMT
        ws.cell(r, c).border = BORDER

    ws.column_dimensions['A'].width = 28
    for c in range(2, 13):
        ws.column_dimensions[get_column_letter(c)].width = 13
    ws.column_dimensions['G'].width = 2


def _write_ccc_sheet(ws, hr, wr):
    ws.merge_cells('A1:L1')
    ws['A1'] = "Cash Conversion Cycle - HCL Technologies vs Wipro"
    ws['A1'].font = Font(name="Calibri", size=12, bold=True, color="1F4E79")
    r = 3
    ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
    ws.cell(r, 2, "HCL Technologies").font = Font(bold=True)
    ws.cell(r, 2).alignment = Alignment(horizontal='center')
    ws.merge_cells(start_row=r, start_column=8, end_row=r, end_column=12)
    ws.cell(r, 8, "Wipro").font = Font(bold=True)
    ws.cell(r, 8).alignment = Alignment(horizontal='center')
    r = 4
    for c, h in enumerate(["Component (days)"] + YEARS + [""] + YEARS, 1):
        ws.cell(r, c, h)
    _style_header(ws, r, 12)

    items = [
        ("Days Inventory Outstanding (DIO)", '_dio'),
        ("Days Sales Outstanding (DSO)", '_dso'),
        ("Days Payable Outstanding (DPO)", '_dpo'),
    ]
    r = 5
    for label, key in items:
        ws.cell(r, 1, label)
        for yi in range(5):
            ws.cell(r, 2+yi, round(hr[key][yi], 1))
            ws.cell(r, 8+yi, round(wr[key][yi], 1))
        _style_data(ws, r, 12, '0.0')
        for c in range(8, 13):
            ws.cell(r, c).number_format = '0.0'
            ws.cell(r, c).border = BORDER
        r += 1

    r += 1
    ws.cell(r, 1, "Cash Conversion Cycle (CCC)")
    ws.cell(r, 1).font = Font(bold=True, size=10, color="1F4E79")
    for yi in range(5):
        ws.cell(r, 2+yi, round(hr['_ccc'][yi], 1))
        ws.cell(r, 8+yi, round(wr['_ccc'][yi], 1))
    _style_data(ws, r, 12, '0.0')
    for c in range(8, 13):
        ws.cell(r, c).number_format = '0.0'
        ws.cell(r, c).border = BORDER

    ws.column_dimensions['A'].width = 38
    for c in range(2, 13):
        ws.column_dimensions[get_column_letter(c)].width = 13
    ws.column_dimensions['G'].width = 2


def create_excel(out_path):
    wb = Workbook()

    ws_cover = wb.active
    ws_cover.title = "Cover"
    ws_cover.merge_cells('A2:F2')
    ws_cover['A2'] = "Financial Statement Analysis"
    ws_cover['A2'].font = Font(size=16, bold=True, color="1F4E79")
    ws_cover['A2'].alignment = Alignment(horizontal='center')
    ws_cover.merge_cells('A3:F3')
    ws_cover['A3'] = "HCL Technologies Ltd. vs Wipro Ltd."
    ws_cover['A3'].font = Font(size=14, bold=True, color="1F4E79")
    ws_cover['A3'].alignment = Alignment(horizontal='center')
    info = [
        (5, "Team Members:", "K P Manoj (2025MB26539)"),
        (6, "", "T Venkata Suresh (2025MB26590)"),
        (7, "", "Japleen Kaur Khorana (2025MB26602)"),
        (8, "", "Varsha PS (2025MB26503)"),
        (9, "", "Prashanth G (2025MB26598)"),
        (10, "Course:", "Financial Management & Accounting"),
        (11, "Programme:", "MBA in AI for Business"),
        (12, "Institution:", "BITS Pilani (WILP)"),
        (13, "Semester:", "1 - Batch Jan 2026"),
        (14, "Period of Analysis:", "FY2021 to FY2025"),
        (15, "Data Source:", "Moneycontrol (Consolidated Statements)"),
    ]
    for r, lbl, val in info:
        if lbl:
            ws_cover.cell(r, 2, lbl).font = Font(bold=True, size=11)
        ws_cover.cell(r, 3, val).font = Font(size=11)
    ws_cover.column_dimensions['B'].width = 22
    ws_cover.column_dimensions['C'].width = 45

    ws_h = wb.create_sheet("HCL_Data")
    hcl_raw = BASE / "HCL Tech-1.xlsx"
    _write_full_data_sheet(ws_h,
                           "HCL Technologies - Consolidated Financial Data (Rs Cr)",
                           hcl_raw)

    ws_w = wb.create_sheet("Wipro_Data")
    wipro_raw = BASE / "Wipro-1.xlsx"
    _write_full_data_sheet(ws_w,
                           "Wipro Ltd. - Consolidated Financial Data (Rs Cr)",
                           wipro_raw)

    ws_rat = wb.create_sheet("Ratios")
    _write_ratios_sheet(ws_rat, HR, WR)

    ws_dup = wb.create_sheet("DuPont")
    _write_dupont_sheet(ws_dup, HR, WR)

    ws_ccc = wb.create_sheet("CCC")
    _write_ccc_sheet(ws_ccc, HR, WR)

    wb.calculation = CalcProperties(calcMode='auto', fullCalcOnLoad=True)
    wb.save(out_path)
    print(f"[OK] Excel saved: {out_path}")


# ==========================  WORD REPORT  ==========================
def _set_doc_defaults(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
    if rFonts is None:
        rPr = style.element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Times New Roman"/>')
        rPr.append(rFonts)

def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(3)
    return h

def _add_para(doc, text, bold=False, italic=False, align=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p

def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def _fmt(v, dp=2):
    if isinstance(v, float):
        return f"{v:,.{dp}f}"
    return str(v)

def create_word_report(out_path):
    doc = Document()
    _set_doc_defaults(doc)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    _add_para(doc, "Financial Statement Analysis", bold=True, size=16,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "HCL Technologies Ltd. vs Wipro Ltd.", bold=True, size=14,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "FY2021 - FY2025 (Consolidated)", bold=True, size=12,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "")
    _add_para(doc, "K P Manoj (2025MB26539)  |  T Venkata Suresh (2025MB26590)  |  Japleen Kaur Khorana (2025MB26602)", bold=True, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "Varsha PS (2025MB26503)  |  Prashanth G (2025MB26598)  |  FMA EC1", bold=True, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "MBA in AI for Business  |  BITS Pilani (WILP)  |  Semester 1, Jan 2026 Batch",
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "")

    # ---- 1. INTRODUCTION ----
    _add_heading(doc, "1. Introduction", level=2)
    _add_para(doc,
        "This report presents a comprehensive financial statement analysis of HCL Technologies Ltd. "
        "(HCLTech) and Wipro Ltd. over the five-year period from FY2021 to FY2025, using consolidated "
        "financial statements sourced from Moneycontrol. The analysis evaluates liquidity, profitability, "
        "solvency, turnover efficiency, and valuation through 19 financial ratios, supplemented by "
        "DuPont decomposition and Cash Conversion Cycle analysis.")
    _add_para(doc,
        "HCL Technologies Ltd., established in 1976 and headquartered in Noida, is one of India's "
        "leading IT services companies with consolidated revenue exceeding Rs 1,17,055 crore in FY2025. "
        "HCLTech operates across 60+ countries, serving clients in banking, financial services, "
        "manufacturing, retail, life sciences, and other sectors. The company's service portfolio spans "
        "IT services, engineering and R&D services, and products & platforms. HCLTech employs over "
        "227,000 professionals globally and has maintained a consistent organic growth strategy, "
        "distinguishing itself from peers who pursued aggressive inorganic expansion.")
    _add_para(doc,
        "Wipro Ltd., founded in 1945 and originally a vegetable oil manufacturer, transformed into an "
        "IT powerhouse under Azim Premji's leadership. With FY2025 consolidated revenue of approximately "
        "Rs 89,088 crore, Wipro is the fourth-largest Indian IT services company. The company serves "
        "clients across BFSI, healthcare, consumer goods, energy, and manufacturing verticals, with "
        "operations spanning 66 countries and a workforce of approximately 230,000. Notably, Wipro "
        "pursued an aggressive acquisition strategy in FY2021-22, including the landmark USD 1.45 billion "
        "acquisition of Capco (a consulting firm specializing in financial services), which significantly "
        "reshaped its revenue profile and balance sheet.")
    _add_para(doc,
        "These two companies are highly comparable: both are top-5 Indian IT services firms competing "
        "for similar clients and talent in the global technology services market. They share similar "
        "business models (primarily services-driven with minimal manufacturing), operate in the same "
        "regulatory environment, and are listed on both NSE and BSE. The comparison is particularly "
        "insightful because it contrasts HCLTech's organic growth model against Wipro's "
        "acquisition-heavy strategy, revealing how different strategic choices manifest in financial "
        "performance over time.")
    _add_para(doc,
        "The analysis period (FY2021-FY2025) is especially significant as it encompasses the COVID-19 "
        "pandemic recovery, a global surge in digital transformation spending, rising interest rates, "
        "and increasing macroeconomic uncertainty. These external factors created both tailwinds and "
        "headwinds for the IT services sector, making it an ideal period to evaluate how each company's "
        "strategy performed under varied conditions. All data has been sourced from Moneycontrol using "
        "consolidated financial statements to ensure comparability and completeness.")

    # ---- 2. INTRA-COMPANY: HCL ----
    _add_heading(doc, "2. Intra-Company Analysis: HCL Technologies", level=2)
    _add_para(doc,
        "Revenue & Growth: HCLTech's consolidated revenue grew from Rs 75,379 crore in FY2021 to "
        "Rs 1,17,055 crore in FY2025, representing a robust 5-year CAGR of 11.6%. Growth was driven "
        "by strong demand for digital transformation services, cloud migration projects, and engineering "
        "R&D services. Revenue acceleration was particularly notable in FY2022 (+13.6%) as post-pandemic "
        "digitization spending surged. Growth moderated to 6.5% in FY2025, reflecting global "
        "macroeconomic caution and delayed deal closures in certain verticals.")
    _add_para(doc,
        "Profitability: The operating profit margin (EBIT/Revenue) ranged between 20.2% and 21.7% over "
        "the five-year period, showing remarkable stability despite industry-wide wage inflation and "
        "supply-side constraints. The net profit margin remained in a tight band of 14.3%-15.8%, with "
        "FY2022 registering the highest margin of 15.78% due to operational efficiencies and favorable "
        "currency movements. Return on Equity (ROE) improved significantly from 18.6% in FY2021 to "
        "25.0% in FY2025. This improvement was driven by consistent profit growth and a controlled "
        "equity base, as HCLTech returned substantial capital through dividends (payout ratio exceeding "
        "85% in recent years), preventing excessive equity buildup. Return on Capital Employed (ROCE) "
        "similarly rose from 23.8% to 30.8%, confirming excellent capital utilization.")

    _add_table(doc,
        ["Ratio", "FY2025", "FY2024", "FY2023", "FY2022", "FY2021"],
        [
            ["Net Profit Margin (%)", _fmt(HR['Net Profit Ratio (%)'][0]), _fmt(HR['Net Profit Ratio (%)'][1]),
             _fmt(HR['Net Profit Ratio (%)'][2]), _fmt(HR['Net Profit Ratio (%)'][3]), _fmt(HR['Net Profit Ratio (%)'][4])],
            ["ROE (%)", _fmt(HR['Return on Equity (%)'][0]), _fmt(HR['Return on Equity (%)'][1]),
             _fmt(HR['Return on Equity (%)'][2]), _fmt(HR['Return on Equity (%)'][3]), _fmt(HR['Return on Equity (%)'][4])],
            ["ROCE (%)", _fmt(HR['Return on Capital Employed (%)'][0]), _fmt(HR['Return on Capital Employed (%)'][1]),
             _fmt(HR['Return on Capital Employed (%)'][2]), _fmt(HR['Return on Capital Employed (%)'][3]), _fmt(HR['Return on Capital Employed (%)'][4])],
            ["Current Ratio", _fmt(HR['Current Ratio'][0]), _fmt(HR['Current Ratio'][1]),
             _fmt(HR['Current Ratio'][2]), _fmt(HR['Current Ratio'][3]), _fmt(HR['Current Ratio'][4])],
            ["Debt-Equity", _fmt(HR['Debt-Equity Ratio'][0]), _fmt(HR['Debt-Equity Ratio'][1]),
             _fmt(HR['Debt-Equity Ratio'][2]), _fmt(HR['Debt-Equity Ratio'][3]), _fmt(HR['Debt-Equity Ratio'][4])],
            ["Interest Coverage", _fmt(HR['Interest Coverage Ratio'][0]), _fmt(HR['Interest Coverage Ratio'][1]),
             _fmt(HR['Interest Coverage Ratio'][2]), _fmt(HR['Interest Coverage Ratio'][3]), _fmt(HR['Interest Coverage Ratio'][4])],
        ])

    _add_para(doc,
        "Liquidity & Solvency: The current ratio remained consistently above 2.2x throughout the "
        "period, peaking at 2.61x in FY2024. The quick ratio was nearly identical (difference <0.01x), "
        "reflecting the asset-light nature of IT services with negligible inventory. Cash and cash "
        "equivalents grew from Rs 8,888 crore to Rs 21,289 crore (+140%), providing a substantial "
        "liquidity buffer. HCLTech maintained an exceptionally conservative capital structure with a "
        "debt-to-equity ratio between 0.03x and 0.06x. Long-term borrowings declined from Rs 3,828 "
        "crore to just Rs 70 crore by FY2025. The interest coverage ratio of 43.5x (FY2025) indicates "
        "virtually zero risk of debt service failure.")
    _add_para(doc,
        "Turnover: The debtors turnover ratio of approximately 4.6x translates to an average collection "
        "period of around 80 days, typical for the IT services industry where large enterprise contracts "
        "often have 60-90 day payment terms. Fixed asset turnover improved from 3.7x to 7.9x, "
        "reflecting the company's increasing shift toward cloud-based delivery models requiring less "
        "physical infrastructure. Earnings per share grew from Rs 41.07 to Rs 64.05, a 56% increase "
        "over the five-year period, rewarding shareholders consistently. Inventory turnover ratios are "
        "extremely high (100x+), which is characteristic of IT services companies that hold minimal "
        "physical inventory. The creditors turnover ratio showed a notable decline in FY2025 (to 2.59x "
        "from 3.78x in FY2024), driven by a spike in trade payables from Rs 5,853 crore to Rs 13,234 "
        "crore, likely reflecting reclassification or timing differences in vendor payments.")

    # ---- 3. INTRA-COMPANY: WIPRO ----
    _add_heading(doc, "3. Intra-Company Analysis: Wipro", level=2)
    _add_para(doc,
        "Revenue & Growth: Wipro's consolidated revenue grew from Rs 61,943 crore in FY2021 to "
        "Rs 89,088 crore in FY2025, a 5-year CAGR of 9.5%. The growth trajectory was uneven: a sharp "
        "27.7% jump in FY2022 was largely attributed to the acquisition of Capco and other entities, "
        "adding significant inorganic revenue. Subsequent years saw growth decelerate markedly, with "
        "FY2024 revenue of Rs 89,760 crore actually declining by 0.8% in FY2025, reflecting challenges "
        "in the European market, leadership transitions, and weak discretionary spending by clients.")
    _add_para(doc,
        "Profitability: The operating profit margin compressed from 23.2% in FY2021 to 21.3% in FY2025, "
        "though the trajectory was non-linear: it dipped to 17.4% in FY2023 before recovering. The "
        "decline was driven by higher subcontracting costs, wage inflation, and integration expenses "
        "from acquisitions. The net profit margin fell from 17.52% in FY2021 to 14.80% in FY2025, "
        "reaching a low of 12.40% in FY2024. ROE declined from 19.66% to 15.94%, partly explained by "
        "a significantly expanded equity base (from Rs 54,901 crore to Rs 82,364 crore) driven by "
        "retained earnings and the Capco acquisition, which has not yet generated proportionate profit "
        "growth. ROCE similarly declined from 24.10% to 19.03%, indicating that incremental capital "
        "deployment yielded diminishing returns.")

    _add_table(doc,
        ["Ratio", "FY2025", "FY2024", "FY2023", "FY2022", "FY2021"],
        [
            ["Net Profit Margin (%)", _fmt(WR['Net Profit Ratio (%)'][0]), _fmt(WR['Net Profit Ratio (%)'][1]),
             _fmt(WR['Net Profit Ratio (%)'][2]), _fmt(WR['Net Profit Ratio (%)'][3]), _fmt(WR['Net Profit Ratio (%)'][4])],
            ["ROE (%)", _fmt(WR['Return on Equity (%)'][0]), _fmt(WR['Return on Equity (%)'][1]),
             _fmt(WR['Return on Equity (%)'][2]), _fmt(WR['Return on Equity (%)'][3]), _fmt(WR['Return on Equity (%)'][4])],
            ["ROCE (%)", _fmt(WR['Return on Capital Employed (%)'][0]), _fmt(WR['Return on Capital Employed (%)'][1]),
             _fmt(WR['Return on Capital Employed (%)'][2]), _fmt(WR['Return on Capital Employed (%)'][3]), _fmt(WR['Return on Capital Employed (%)'][4])],
            ["Current Ratio", _fmt(WR['Current Ratio'][0]), _fmt(WR['Current Ratio'][1]),
             _fmt(WR['Current Ratio'][2]), _fmt(WR['Current Ratio'][3]), _fmt(WR['Current Ratio'][4])],
            ["Debt-Equity", _fmt(WR['Debt-Equity Ratio'][0]), _fmt(WR['Debt-Equity Ratio'][1]),
             _fmt(WR['Debt-Equity Ratio'][2]), _fmt(WR['Debt-Equity Ratio'][3]), _fmt(WR['Debt-Equity Ratio'][4])],
            ["Interest Coverage", _fmt(WR['Interest Coverage Ratio'][0]), _fmt(WR['Interest Coverage Ratio'][1]),
             _fmt(WR['Interest Coverage Ratio'][2]), _fmt(WR['Interest Coverage Ratio'][3]), _fmt(WR['Interest Coverage Ratio'][4])],
        ])

    _add_para(doc,
        "Liquidity: Wipro's current ratio improved from 2.27x in FY2021 to 2.72x in FY2025, "
        "supported by a large portfolio of current investments (Rs 41,147 crore, nearly 4x HCLTech's). "
        "Cash holdings stood at Rs 12,197 crore. The absolute liquid ratio of 1.86x provides strong "
        "short-term coverage.")
    _add_para(doc,
        "Solvency: The debt-to-equity ratio increased from 0.12x in FY2021 to 0.20x in FY2025, "
        "reflecting debt raised to finance the Capco acquisition (long-term borrowings rose from "
        "Rs 746 crore to Rs 6,395 crore). While still manageable, this represents a meaningful shift "
        "in capital structure. Interest coverage declined from 33.73x to 14.83x, consistent with "
        "higher debt levels and elevated interest rates in the post-pandemic period.")
    _add_para(doc,
        "Turnover & Efficiency: The debtors turnover ratio exhibited significant volatility, with trade "
        "receivables spiking to Rs 18,687 crore in FY2023 (likely due to Capco integration effects and "
        "changed billing cycles) before normalizing to Rs 11,775 crore in FY2025. This resulted in "
        "average collection periods ranging from 48 days (FY2022) to 73 days (FY2024). Inventory "
        "metrics are similarly characteristic of IT services, with near-zero holding days. The fixed "
        "asset turnover of 6.71x in FY2025 reflects efficient infrastructure utilization. EPS declined "
        "from Rs 19.71 (FY2021) to Rs 12.54 (FY2025), though this largely reflects the 1:1 bonus issue "
        "in FY2025 that doubled the share count from 522.5 crore to 1,047.2 crore shares, rather than "
        "a deterioration in profitability.")

    # ---- 4. INTER-COMPANY ----
    _add_heading(doc, "4. Inter-Company Comparison", level=2)
    _add_para(doc,
        "Using FY2025 data as the primary reference point and 5-year trends for context, this section "
        "directly compares HCLTech and Wipro across the key ratio categories.")

    _add_table(doc,
        ["Ratio", "HCL FY2025", "Wipro FY2025", "Better"],
        [
            ["Revenue (Rs Cr)", "1,17,055", "89,088", "HCL"],
            ["5-Year Revenue CAGR", "11.6%", "9.5%", "HCL"],
            ["Net Profit Margin (%)", _fmt(HR['Net Profit Ratio (%)'][0]), _fmt(WR['Net Profit Ratio (%)'][0]), "HCL"],
            ["ROE (%)", _fmt(HR['Return on Equity (%)'][0]), _fmt(WR['Return on Equity (%)'][0]), "HCL"],
            ["ROCE (%)", _fmt(HR['Return on Capital Employed (%)'][0]), _fmt(WR['Return on Capital Employed (%)'][0]), "HCL"],
            ["Current Ratio", _fmt(HR['Current Ratio'][0]), _fmt(WR['Current Ratio'][0]), "Wipro"],
            ["Debt-Equity Ratio", _fmt(HR['Debt-Equity Ratio'][0]), _fmt(WR['Debt-Equity Ratio'][0]), "HCL"],
            ["Interest Coverage", _fmt(HR['Interest Coverage Ratio'][0],1), _fmt(WR['Interest Coverage Ratio'][0],1), "HCL"],
            ["EPS (Rs)", _fmt(HR['Earnings per Share (Rs)'][0]), _fmt(WR['Earnings per Share (Rs)'][0]), "HCL"],
            ["P/E Ratio", _fmt(HR['P/E Ratio'][0]), _fmt(WR['P/E Ratio'][0]), "-"],
        ])

    _add_para(doc,
        "Scale & Growth: HCLTech's FY2025 revenue of Rs 1,17,055 crore is 31% larger than Wipro's "
        "Rs 89,088 crore. More importantly, HCLTech's 5-year revenue CAGR of 11.6% surpasses Wipro's "
        "9.5%, despite the latter's aggressive acquisition program. This suggests that HCLTech's organic "
        "growth strategy has been more effective at scaling revenue sustainably than Wipro's inorganic "
        "approach.")
    _add_para(doc,
        "Profitability: While net profit margins are comparable in FY2025 (HCL 14.86% vs Wipro 14.80%), "
        "HCLTech has been consistently more profitable over the 5-year period. The divergence is starker "
        "in return metrics: ROE of 25.0% vs 15.9% and ROCE of 30.8% vs 19.0% demonstrate that HCLTech "
        "generates far superior returns on invested capital. This gap is primarily driven by HCLTech's "
        "higher asset turnover efficiency, as the DuPont analysis in Section 5 confirms.")
    _add_para(doc,
        "Solvency: HCLTech's near-zero debt (D/E 0.03x) contrasts sharply with Wipro's 0.20x, which "
        "carries Rs 16,182 crore in total debt versus HCLTech's Rs 2,291 crore. The interest coverage "
        "gap is even more telling: 43.5x vs 14.8x. HCLTech's conservative financial structure provides "
        "greater strategic flexibility and lower financial risk.")
    _add_para(doc,
        "Liquidity & Turnover: Wipro holds a marginal advantage in current ratio (2.72x vs 2.22x), "
        "primarily due to its larger current investment portfolio (Rs 41,147 crore vs Rs 7,473 crore). "
        "Both companies display the turnover characteristics typical of IT services: near-zero "
        "inventory, receivables collection periods of 60-80 days, and high fixed asset turnover. "
        "Wipro's FY2025 debtor turnover of 6.11x (vs HCL's 4.56x) appears stronger, but this is "
        "partly due to normalization of previously elevated receivables rather than structural "
        "efficiency improvement.")
    _add_para(doc,
        "Overall Assessment: Across the majority of financial parameters - revenue growth, "
        "profitability margins, return on equity and capital, solvency, and interest coverage - "
        "HCLTech demonstrates clear and consistent superiority. Wipro's only notable advantages are "
        "its higher current ratio and faster debtor collection in FY2025. The strategic implications "
        "are significant: HCLTech's organic growth model has delivered more predictable, higher-quality "
        "returns, while Wipro's acquisition-heavy approach has yet to translate into proportionate "
        "financial outperformance.")

    # ---- 5. VALUE-ADDED + CONCLUSION ----
    _add_heading(doc, "5. Value-Added Analysis & Conclusion", level=2)

    _add_para(doc, "5.1 DuPont Analysis", bold=True)
    _add_para(doc,
        "The DuPont framework decomposes Return on Equity into three multiplicative drivers: "
        "Net Profit Margin (profitability), Asset Turnover (efficiency), and Equity Multiplier "
        "(leverage). This decomposition reveals the underlying drivers of ROE differences between "
        "the two companies.")

    _add_table(doc,
        ["DuPont Component", "HCL FY2025", "HCL FY2021", "Wipro FY2025", "Wipro FY2021"],
        [
            ["Net Profit Margin", f"{HR['_npm'][0]:.4f}", f"{HR['_npm'][4]:.4f}",
             f"{WR['_npm'][0]:.4f}", f"{WR['_npm'][4]:.4f}"],
            ["Asset Turnover", f"{HR['_at'][0]:.4f}", f"{HR['_at'][4]:.4f}",
             f"{WR['_at'][0]:.4f}", f"{WR['_at'][4]:.4f}"],
            ["Equity Multiplier", f"{HR['_em'][0]:.4f}", f"{HR['_em'][4]:.4f}",
             f"{WR['_em'][0]:.4f}", f"{WR['_em'][4]:.4f}"],
            ["ROE (%)", f"{HR['_roe_d'][0]:.2f}", f"{HR['_roe_d'][4]:.2f}",
             f"{WR['_roe_d'][0]:.2f}", f"{WR['_roe_d'][4]:.2f}"],
        ])

    _add_para(doc,
        "DuPont Derivation - HCL FY2025:")
    _add_para(doc,
        f"  Net Profit Margin = PAT / Revenue = 17,399 / 1,17,055 = {HR['_npm'][0]:.4f} "
        f"(PAT from HCL P&L; Revenue from HCL P&L)", size=10)
    _add_para(doc,
        f"  Asset Turnover = Revenue / Total Assets = 1,17,055 / 1,05,544 = {HR['_at'][0]:.4f} "
        f"(Revenue from HCL P&L; Total Assets from HCL Balance Sheet)", size=10)
    _add_para(doc,
        f"  Equity Multiplier = Total Assets / Total Shareholders' Funds = 1,05,544 / 69,655 = {HR['_em'][0]:.4f} "
        f"(both from HCL Balance Sheet)", size=10)
    _add_para(doc,
        f"  ROE = {HR['_npm'][0]:.4f} x {HR['_at'][0]:.4f} x {HR['_em'][0]:.4f} x 100 = "
        f"{HR['_roe_d'][0]:.2f}%", size=10)
    _add_para(doc,
        "DuPont Derivation - Wipro FY2025:")
    _add_para(doc,
        f"  Net Profit Margin = PAT / Revenue = 13,192.6 / 89,088.4 = {WR['_npm'][0]:.4f} "
        f"(PAT from Wipro P&L; Revenue from Wipro P&L)", size=10)
    _add_para(doc,
        f"  Asset Turnover = Revenue / Total Assets = 89,088.4 / 1,28,185.2 = {WR['_at'][0]:.4f} "
        f"(Revenue from Wipro P&L; Total Assets from Wipro Balance Sheet)", size=10)
    _add_para(doc,
        f"  Equity Multiplier = Total Assets / Total Shareholders' Funds = 1,28,185.2 / 82,364.1 = {WR['_em'][0]:.4f} "
        f"(both from Wipro Balance Sheet)", size=10)
    _add_para(doc,
        f"  ROE = {WR['_npm'][0]:.4f} x {WR['_at'][0]:.4f} x {WR['_em'][0]:.4f} x 100 = "
        f"{WR['_roe_d'][0]:.2f}%", size=10)

    _add_para(doc,
        "The DuPont analysis reveals that the 9-percentage-point gap in ROE between HCLTech (25.0%) "
        "and Wipro (16.0%) is almost entirely attributable to asset turnover (1.109 vs 0.695). Both "
        "companies have similar profit margins (~14.8%) and comparable equity multipliers (1.52 vs "
        "1.56). The lower asset turnover for Wipro reflects the significant goodwill and intangible "
        "assets from the Capco acquisition (intangibles of Rs 2,745 crore), which inflate the asset "
        "base without generating proportional revenue. Over 5 years, HCLTech's asset turnover improved "
        "from 0.87x to 1.11x, while Wipro's deteriorated from 0.75x to 0.70x, confirming that "
        "acquisitions expanded Wipro's balance sheet faster than its revenue.")

    _add_para(doc, "5.2 Cash Conversion Cycle (CCC)", bold=True)
    _add_para(doc,
        "The Cash Conversion Cycle measures the time (in days) it takes for a company to convert its "
        "investments in inventory and other resources into cash flows from sales. "
        "CCC = Days Inventory Outstanding (DIO) + Days Sales Outstanding (DSO) - Days Payable "
        "Outstanding (DPO).")

    _add_table(doc,
        ["CCC Component", "HCL FY2025", "HCL FY2021", "Wipro FY2025", "Wipro FY2021"],
        [
            ["DIO (days)", f"{HR['_dio'][0]:.1f}", f"{HR['_dio'][4]:.1f}",
             f"{WR['_dio'][0]:.1f}", f"{WR['_dio'][4]:.1f}"],
            ["DSO (days)", f"{HR['_dso'][0]:.1f}", f"{HR['_dso'][4]:.1f}",
             f"{WR['_dso'][0]:.1f}", f"{WR['_dso'][4]:.1f}"],
            ["DPO (days)", f"{HR['_dpo'][0]:.1f}", f"{HR['_dpo'][4]:.1f}",
             f"{WR['_dpo'][0]:.1f}", f"{WR['_dpo'][4]:.1f}"],
            ["CCC (days)", f"{HR['_ccc'][0]:.1f}", f"{HR['_ccc'][4]:.1f}",
             f"{WR['_ccc'][0]:.1f}", f"{WR['_ccc'][4]:.1f}"],
        ])

    h_cogs_v = H_PNL['purchase'][0] + H_PNL['op_exp'][0] + H_PNL['inv_change'][0]
    w_cogs_v = W_PNL['purchase'][0] + W_PNL['op_exp'][0] + W_PNL['inv_change'][0]
    h_purch = H_PNL['purchase'][0] + H_PNL['op_exp'][0] + H_PNL['other_exp'][0]
    w_purch = W_PNL['purchase'][0] + W_PNL['op_exp'][0] + W_PNL['other_exp'][0]

    _add_para(doc, "CCC Derivation - HCL FY2025:", bold=False, size=10)
    _add_para(doc,
        f"  DIO = 365 x Avg(Inventories FY25, FY24) / COGS = 365 x Avg(133, 185) / {h_cogs_v:,} = {HR['_dio'][0]:.1f} days "
        f"(Inventories from HCL Balance Sheet; COGS = Purchase + Operating Exp + Inv Changes from HCL P&L)", size=10)
    _add_para(doc,
        f"  DSO = 365 x Avg(Trade Rec FY25, FY24) / Revenue = 365 x Avg(25,842, 25,521) / 1,17,055 = {HR['_dso'][0]:.1f} days "
        f"(Trade Receivables from HCL Balance Sheet; Revenue from HCL P&L)", size=10)
    _add_para(doc,
        f"  DPO = 365 x Avg(Trade Pay FY25, FY24) / Purchases = 365 x Avg(13,234, 5,853) / {h_purch:,} = {HR['_dpo'][0]:.1f} days "
        f"(Trade Payables from HCL Balance Sheet; Purchases = Purchase + Op Exp + Other Exp from HCL P&L)", size=10)
    _add_para(doc,
        f"  CCC = DIO + DSO - DPO = {HR['_dio'][0]:.1f} + {HR['_dso'][0]:.1f} - {HR['_dpo'][0]:.1f} = {HR['_ccc'][0]:.1f} days", size=10)

    _add_para(doc, "CCC Derivation - Wipro FY2025:", bold=False, size=10)
    _add_para(doc,
        f"  DIO = 365 x Avg(69.4, 90.7) / {w_cogs_v:,.1f} = {WR['_dio'][0]:.1f} days "
        f"(Inventories from Wipro Balance Sheet; COGS from Wipro P&L)", size=10)
    _add_para(doc,
        f"  DSO = 365 x Avg(11,774.5, 17,382.2) / 89,088.4 = {WR['_dso'][0]:.1f} days "
        f"(Trade Receivables from Wipro Balance Sheet; Revenue from Wipro P&L)", size=10)
    _add_para(doc,
        f"  DPO = 365 x Avg(5,866.7, 5,765.5) / {w_purch:,.1f} = {WR['_dpo'][0]:.1f} days "
        f"(Trade Payables from Wipro Balance Sheet; Purchases from Wipro P&L)", size=10)
    _add_para(doc,
        f"  CCC = {WR['_dio'][0]:.1f} + {WR['_dso'][0]:.1f} - {WR['_dpo'][0]:.1f} = {WR['_ccc'][0]:.1f} days", size=10)

    _add_para(doc,
        f"Both companies exhibit negative CCCs (HCL: {HR['_ccc'][0]:.1f} days, "
        f"Wipro: {WR['_ccc'][0]:.1f} days), which is typical of IT services firms that collect "
        "from clients relatively quickly compared to their payment obligations, and hold negligible "
        "inventory. HCLTech's CCC is more negative, indicating superior working capital efficiency. "
        "However, HCLTech's DPO of 203 days is elevated in FY2025 due to a spike in trade payables "
        "(from Rs 5,853 crore to Rs 13,234 crore), which may reflect reclassification between "
        "'Other Current Liabilities' and 'Trade Payables'. Both firms benefit from the IT services "
        "model where human capital rather than physical inventory drives operations, keeping DIO "
        "near zero across all years.")

    _add_para(doc, "5.3 Conclusion", bold=True)
    _add_para(doc,
        "Based on this comprehensive analysis of 19 financial ratios, DuPont decomposition, and "
        "Cash Conversion Cycle over FY2021-FY2025, HCL Technologies Ltd. emerges as the clearly "
        "superior performer compared to Wipro Ltd. HCLTech's key advantages include:")
    advantages = [
        "Stronger organic revenue growth (CAGR 11.6% vs 9.5%) without acquisition-related risks",
        "Significantly higher return on equity (25.0% vs 15.9%) and return on capital employed (30.8% vs 19.0%)",
        "Near debt-free balance sheet (D/E 0.03x vs 0.20x) providing strategic flexibility",
        "Superior DuPont asset turnover (1.11x vs 0.70x), confirming efficient capital utilization",
        "More negative Cash Conversion Cycle, reflecting better working capital management",
        "Consistent EPS growth of 56% over 5 years, directly benefiting shareholders",
    ]
    for adv in advantages:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(adv)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)

    _add_para(doc,
        "Wipro's strengths - higher current ratio and larger investment portfolio - are overshadowed by "
        "its declining profitability trajectory, leveraged balance sheet, and failure to generate "
        "proportionate returns from its acquisition-driven growth strategy. The Capco acquisition, "
        "while strategically sound for expanding financial services consulting capabilities, has diluted "
        "returns and inflated the asset base. Until Wipro demonstrates sustainable margin improvement "
        "and asset turnover recovery, HCLTech remains the better performer from a financial analysis "
        "perspective. The market recognizes this: HCLTech's P/E ratio of {:.1f}x versus Wipro's "
        "{:.1f}x reflects the premium investors are willing to pay for HCLTech's superior and more "
        "consistent financial performance.".format(HR['P/E Ratio'][0], WR['P/E Ratio'][0]))

    # ---- REFERENCES ----
    _add_heading(doc, "References", level=2)
    refs = [
        "HCL Technologies Ltd. - Consolidated Financial Statements (FY2021-FY2025), Moneycontrol.",
        "Wipro Ltd. - Consolidated Financial Statements (FY2021-FY2025), Moneycontrol.",
        "HCL Technologies Ltd. - Key Financial Ratios, Moneycontrol.",
        "Wipro Ltd. - Key Financial Ratios, Moneycontrol.",
        "HCL Technologies Annual Reports (FY2021-FY2025).",
        "Wipro Ltd. Annual Reports (FY2021-FY2025).",
    ]
    for ref in refs:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)

    doc.save(out_path)
    print(f"[OK] Word report saved: {out_path}")


# ==========================  RATIO PROOFS (Enhanced with Source Derivations)  ==========================
def create_ratio_proofs(out_path):
    doc = Document()
    _set_doc_defaults(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    _add_para(doc, "Ratio Calculation Proofs", bold=True, size=14,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "HCL Technologies vs Wipro | FY2025 Detailed Calculations", size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "K P Manoj (2025MB26539) | T Venkata Suresh (2025MB26590) | Japleen Kaur Khorana (2025MB26602)", size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "Varsha PS (2025MB26503) | Prashanth G (2025MB26598) | FMA EC1", size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "Data Source: HCL Tech-1.xlsx and Wipro-1.xlsx (Moneycontrol Consolidated)", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "All values in Rs. Crore unless stated otherwise.", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    _add_para(doc, "")

    def proof_enhanced(name, formula, hcl_inputs, hcl_calc, hcl_result,
                       wip_inputs, wip_calc, wip_result, interpretation):
        """Enhanced proof with source derivation for each input value."""
        _add_para(doc, name, bold=True, size=11)
        _add_para(doc, f"Formula: {formula}", italic=True, size=10)
        _add_para(doc, "")

        _add_para(doc, "HCL Technologies FY2025 - Input Values:", bold=True, size=10)
        for inp in hcl_inputs:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(inp)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(1)
        _add_para(doc, f"Substitution: {hcl_calc}", size=10)
        _add_para(doc, f"Result: {hcl_result}", size=10, bold=True)
        _add_para(doc, "")

        _add_para(doc, "Wipro FY2025 - Input Values:", bold=True, size=10)
        for inp in wip_inputs:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(inp)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(1)
        _add_para(doc, f"Substitution: {wip_calc}", size=10)
        _add_para(doc, f"Result: {wip_result}", size=10, bold=True)
        _add_para(doc, "")

        _add_para(doc, f"Interpretation: {interpretation}", size=10, italic=True)
        _add_para(doc, "")

    # ===================== A. LIQUIDITY RATIOS =====================
    _add_heading(doc, "A. Liquidity Ratios", level=2)

    proof_enhanced(
        "1. Current Ratio",
        "Total Current Assets / Total Current Liabilities",
        [
            "Total Current Assets = 62,109 (from HCL Balance Sheet \u2192 'Total Current Assets' row)",
            "Total Current Liabilities = 28,039 (from HCL Balance Sheet \u2192 'Total Current Liabilities' row)",
        ],
        "Current Ratio = 62,109 / 28,039",
        f"Current Ratio = {HR['Current Ratio'][0]:.2f}x",
        [
            "Total Current Assets = 77,777.5 (from Wipro Balance Sheet \u2192 'Total Current Assets' row)",
            "Total Current Liabilities = 28,625.3 (from Wipro Balance Sheet \u2192 'Total Current Liabilities' row)",
        ],
        "Current Ratio = 77,777.5 / 28,625.3",
        f"Current Ratio = {WR['Current Ratio'][0]:.2f}x",
        "Both companies maintain healthy liquidity (>2x). Wipro has a marginally higher ratio "
        "due to its large current investment portfolio (Rs 41,147 Cr vs HCL's Rs 7,473 Cr)."
    )

    proof_enhanced(
        "2. Quick Ratio",
        "(Total Current Assets - Inventories) / Total Current Liabilities",
        [
            "Total Current Assets = 62,109 (from HCL Balance Sheet \u2192 'Total Current Assets' row)",
            "Inventories = 133 (from HCL Balance Sheet \u2192 'Inventories' row)",
            "Total Current Liabilities = 28,039 (from HCL Balance Sheet \u2192 'Total Current Liabilities' row)",
        ],
        "Quick Ratio = (62,109 - 133) / 28,039 = 61,976 / 28,039",
        f"Quick Ratio = {HR['Quick Ratio'][0]:.2f}x",
        [
            "Total Current Assets = 77,777.5 (from Wipro Balance Sheet \u2192 'Total Current Assets' row)",
            "Inventories = 69.4 (from Wipro Balance Sheet \u2192 'Inventories' row)",
            "Total Current Liabilities = 28,625.3 (from Wipro Balance Sheet \u2192 'Total Current Liabilities' row)",
        ],
        "Quick Ratio = (77,777.5 - 69.4) / 28,625.3 = 77,708.1 / 28,625.3",
        f"Quick Ratio = {WR['Quick Ratio'][0]:.2f}x",
        "Quick ratio is nearly identical to the current ratio for both companies because "
        "IT services firms hold negligible inventory (HCL: Rs 133 Cr, Wipro: Rs 69 Cr)."
    )

    proof_enhanced(
        "3. Absolute Liquid Ratio",
        "(Cash and Cash Equivalents + Current Investments) / Total Current Liabilities",
        [
            "Cash and Cash Equivalents = 21,289 (from HCL Balance Sheet \u2192 'Cash And Cash Equivalents' row)",
            "Current Investments = 7,473 (from HCL Balance Sheet \u2192 'Current Investments' row)",
            "Total Current Liabilities = 28,039 (from HCL Balance Sheet \u2192 'Total Current Liabilities' row)",
        ],
        "Absolute Liquid Ratio = (21,289 + 7,473) / 28,039 = 28,762 / 28,039",
        f"Absolute Liquid Ratio = {HR['Absolute Liquid Ratio'][0]:.2f}x",
        [
            "Cash and Cash Equivalents = 12,197.4 (from Wipro Balance Sheet \u2192 'Cash And Cash Equivalents' row)",
            "Current Investments = 41,147.4 (from Wipro Balance Sheet \u2192 'Current Investments' row)",
            "Total Current Liabilities = 28,625.3 (from Wipro Balance Sheet \u2192 'Total Current Liabilities' row)",
        ],
        "Absolute Liquid Ratio = (12,197.4 + 41,147.4) / 28,625.3 = 53,344.8 / 28,625.3",
        f"Absolute Liquid Ratio = {WR['Absolute Liquid Ratio'][0]:.2f}x",
        "Wipro's absolute liquid ratio is significantly higher (1.86x vs 1.03x) due to its "
        "massive current investment portfolio. Both are well above the standard benchmark of 0.5x."
    )

    # ===================== B. PROFITABILITY RATIOS =====================
    _add_heading(doc, "B. Profitability Ratios", level=2)

    h_cogs = H_PNL['purchase'][0] + H_PNL['op_exp'][0] + H_PNL['inv_change'][0]
    w_cogs = W_PNL['purchase'][0] + W_PNL['op_exp'][0] + W_PNL['inv_change'][0]

    proof_enhanced(
        "4. Gross Profit Ratio (%)",
        "(Revenue - COGS) / Revenue x 100\nwhere COGS = Purchase of Stock-in-Trade + Operating & Direct Expenses + Changes in Inventories",
        [
            "Revenue = 1,17,055 (from HCL P&L \u2192 'Revenue From Operations [Net]' row)",
            f"Purchase of Stock-in-Trade = {H_PNL['purchase'][0]:,} (from HCL P&L \u2192 'Purchase Of Stock-In Trade' row)",
            f"Operating & Direct Expenses = {H_PNL['op_exp'][0]:,} (from HCL P&L \u2192 'Operating And Direct Expenses' row)",
            f"Changes in Inventories = {H_PNL['inv_change'][0]} (from HCL P&L \u2192 'Changes In Inventories' row)",
            f"COGS = {H_PNL['purchase'][0]:,} + {H_PNL['op_exp'][0]:,} + {H_PNL['inv_change'][0]} = {h_cogs:,}",
        ],
        f"Gross Profit Ratio = (1,17,055 - {h_cogs:,}) / 1,17,055 x 100 = {117055 - h_cogs:,} / 1,17,055 x 100",
        f"Gross Profit Ratio = {HR['Gross Profit Ratio (%)'][0]:.2f}%",
        [
            "Revenue = 89,088.4 (from Wipro P&L \u2192 'Revenue From Operations [Net]' row)",
            f"Purchase of Stock-in-Trade = {W_PNL['purchase'][0]} (from Wipro P&L \u2192 'Purchase Of Stock-In Trade' row)",
            f"Operating & Direct Expenses = {W_PNL['op_exp'][0]:,} (from Wipro P&L \u2192 'Operating And Direct Expenses' row)",
            f"Changes in Inventories = {W_PNL['inv_change'][0]} (from Wipro P&L \u2192 'Changes In Inventories' row)",
            f"COGS = {W_PNL['purchase'][0]} + {W_PNL['op_exp'][0]:,} + {W_PNL['inv_change'][0]} = {w_cogs:,.1f}",
        ],
        f"Gross Profit Ratio = (89,088.4 - {w_cogs:,.1f}) / 89,088.4 x 100",
        f"Gross Profit Ratio = {WR['Gross Profit Ratio (%)'][0]:.2f}%",
        "High GP ratios (>85%) are normal for IT services where direct material costs are minimal. "
        "The primary cost driver is employee expenses, which is classified separately."
    )

    h_ebit = H_PNL['pbt'][0] + H_PNL['fin_cost'][0]
    w_ebit = W_PNL['pbt'][0] + W_PNL['fin_cost'][0]

    proof_enhanced(
        "5. Operating Profit Ratio (%) - EBIT Margin",
        "EBIT / Revenue x 100\nwhere EBIT = Profit Before Tax (PBT) + Finance Costs",
        [
            f"PBT = {H_PNL['pbt'][0]:,} (from HCL P&L \u2192 'Profit/Loss Before Tax' row)",
            f"Finance Costs = {H_PNL['fin_cost'][0]} (from HCL P&L \u2192 'Finance Costs' row)",
            f"EBIT = {H_PNL['pbt'][0]:,} + {H_PNL['fin_cost'][0]} = {h_ebit:,}",
            "Revenue = 1,17,055 (from HCL P&L \u2192 'Revenue From Operations [Net]' row)",
        ],
        f"Operating Profit Ratio = {h_ebit:,} / 1,17,055 x 100",
        f"Operating Profit Ratio = {HR['Operating Profit Ratio (%)'][0]:.2f}%",
        [
            f"PBT = {W_PNL['pbt'][0]:,.1f} (from Wipro P&L \u2192 'Profit/Loss Before Tax' row)",
            f"Finance Costs = {W_PNL['fin_cost'][0]:,} (from Wipro P&L \u2192 'Finance Costs' row)",
            f"EBIT = {W_PNL['pbt'][0]:,.1f} + {W_PNL['fin_cost'][0]:,} = {w_ebit:,.1f}",
            "Revenue = 89,088.4 (from Wipro P&L \u2192 'Revenue From Operations [Net]' row)",
        ],
        f"Operating Profit Ratio = {w_ebit:,.1f} / 89,088.4 x 100",
        f"Operating Profit Ratio = {WR['Operating Profit Ratio (%)'][0]:.2f}%",
        "HCL's EBIT margin (20.42%) marginally trails Wipro (21.26%) in FY2025. However, HCL has "
        "been more consistent over 5 years with less volatility in operating margins."
    )

    proof_enhanced(
        "6. Net Profit Ratio (%)",
        "Profit After Tax (PAT) / Revenue x 100",
        [
            f"PAT = {H_PNL['pat'][0]:,} (from HCL P&L \u2192 'Profit/Loss After Tax' row)",
            "Revenue = 1,17,055 (from HCL P&L \u2192 'Revenue From Operations [Net]' row)",
        ],
        f"Net Profit Ratio = {H_PNL['pat'][0]:,} / 1,17,055 x 100",
        f"Net Profit Ratio = {HR['Net Profit Ratio (%)'][0]:.2f}%",
        [
            f"PAT = {W_PNL['pat'][0]:,.1f} (from Wipro P&L \u2192 'Profit/Loss After Tax' row)",
            "Revenue = 89,088.4 (from Wipro P&L \u2192 'Revenue From Operations [Net]' row)",
        ],
        f"Net Profit Ratio = {W_PNL['pat'][0]:,.1f} / 89,088.4 x 100",
        f"Net Profit Ratio = {WR['Net Profit Ratio (%)'][0]:.2f}%",
        "Both companies have comparable NPM in FY2025 (~14.8%), but HCL has maintained this level "
        "more consistently while Wipro dipped to 12.4% in FY2024 before recovering."
    )

    h_avg_eq = (H_BS['total_shf'][0] + H_BS['total_shf'][1]) / 2
    w_avg_eq = (W_BS['total_shf'][0] + W_BS['total_shf'][1]) / 2

    proof_enhanced(
        "7. Return on Equity (ROE) (%)",
        "PAT / Average Shareholders' Funds x 100\nAverage = (Opening + Closing) / 2",
        [
            f"PAT = {H_PNL['pat'][0]:,} (from HCL P&L \u2192 'Profit/Loss After Tax' row)",
            f"Total Shareholders' Funds FY2025 = {H_BS['total_shf'][0]:,} (from HCL Balance Sheet \u2192 'Total Shareholders Funds' row)",
            f"Total Shareholders' Funds FY2024 = {H_BS['total_shf'][1]:,} (from HCL Balance Sheet FY2024 \u2192 same row)",
            f"Average Shareholders' Funds = ({H_BS['total_shf'][0]:,} + {H_BS['total_shf'][1]:,}) / 2 = {h_avg_eq:,.0f}",
        ],
        f"ROE = {H_PNL['pat'][0]:,} / {h_avg_eq:,.0f} x 100",
        f"ROE = {HR['Return on Equity (%)'][0]:.2f}%",
        [
            f"PAT = {W_PNL['pat'][0]:,.1f} (from Wipro P&L \u2192 'Profit/Loss After Tax' row)",
            f"Total Shareholders' Funds FY2025 = {W_BS['total_shf'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Total Shareholders Funds' row)",
            f"Total Shareholders' Funds FY2024 = {W_BS['total_shf'][1]:,} (from Wipro Balance Sheet FY2024 \u2192 same row)",
            f"Average Shareholders' Funds = ({W_BS['total_shf'][0]:,.1f} + {W_BS['total_shf'][1]:,}) / 2 = {w_avg_eq:,.1f}",
        ],
        f"ROE = {W_PNL['pat'][0]:,.1f} / {w_avg_eq:,.1f} x 100",
        f"ROE = {WR['Return on Equity (%)'][0]:.2f}%",
        "HCL's ROE (25.24%) is significantly higher than Wipro's (16.82%), indicating HCL generates "
        "substantially more profit per rupee of shareholders' equity."
    )

    h_ce = H_BS['total_assets'][0] - H_BS['total_cl'][0]
    w_ce = W_BS['total_assets'][0] - W_BS['total_cl'][0]

    proof_enhanced(
        "8. Return on Capital Employed (ROCE) (%)",
        "EBIT / Capital Employed x 100\nwhere Capital Employed = Total Assets - Total Current Liabilities",
        [
            f"EBIT = {h_ebit:,} (derived: PBT {H_PNL['pbt'][0]:,} + Finance Costs {H_PNL['fin_cost'][0]} from HCL P&L)",
            f"Total Assets = {H_BS['total_assets'][0]:,} (from HCL Balance Sheet \u2192 'Total Assets' row)",
            f"Total Current Liabilities = {H_BS['total_cl'][0]:,} (from HCL Balance Sheet \u2192 'Total Current Liabilities' row)",
            f"Capital Employed = {H_BS['total_assets'][0]:,} - {H_BS['total_cl'][0]:,} = {h_ce:,}",
        ],
        f"ROCE = {h_ebit:,} / {h_ce:,} x 100",
        f"ROCE = {HR['Return on Capital Employed (%)'][0]:.2f}%",
        [
            f"EBIT = {w_ebit:,.1f} (derived: PBT {W_PNL['pbt'][0]:,.1f} + Finance Costs {W_PNL['fin_cost'][0]:,} from Wipro P&L)",
            f"Total Assets = {W_BS['total_assets'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Total Assets' row)",
            f"Total Current Liabilities = {W_BS['total_cl'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Total Current Liabilities' row)",
            f"Capital Employed = {W_BS['total_assets'][0]:,.1f} - {W_BS['total_cl'][0]:,.1f} = {w_ce:,.1f}",
        ],
        f"ROCE = {w_ebit:,.1f} / {w_ce:,.1f} x 100",
        f"ROCE = {WR['Return on Capital Employed (%)'][0]:.2f}%",
        "HCL's ROCE (30.84%) far exceeds Wipro's (19.01%), confirming that HCL utilizes its total "
        "long-term capital more efficiently to generate operating profits."
    )

    h_sh = H_BS['eq_capital'][0] / FACE_VALUE
    w_sh = W_BS['eq_capital'][0] / FACE_VALUE

    proof_enhanced(
        "9. Earnings per Share (EPS)",
        "PAT After Minority Interest / Number of Shares\nShares Outstanding = Equity Share Capital / Face Value (Rs 2 per share)",
        [
            f"PAT After MI = {H_PNL['pat_mi'][0]:,} (from HCL P&L \u2192 'Consolidated Profit/Loss After MI' row)",
            f"Equity Share Capital = {H_BS['eq_capital'][0]} (from HCL Balance Sheet \u2192 'Equity Share Capital' row)",
            f"Face Value = Rs {FACE_VALUE} per share",
            f"Shares Outstanding = {H_BS['eq_capital'][0]} / {FACE_VALUE} = {h_sh:.1f} Cr shares",
        ],
        f"EPS = {H_PNL['pat_mi'][0]:,} / {h_sh:.1f} = Rs {HR['Earnings per Share (Rs)'][0]:.2f}",
        f"EPS = Rs {HR['Earnings per Share (Rs)'][0]:.2f} per share",
        [
            f"PAT After MI = {W_PNL['pat_mi'][0]:,.1f} (from Wipro P&L \u2192 'Consolidated Profit/Loss After MI' row)",
            f"Equity Share Capital = {W_BS['eq_capital'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Equity Share Capital' row)",
            f"Face Value = Rs {FACE_VALUE} per share",
            f"Shares Outstanding = {W_BS['eq_capital'][0]:,.1f} / {FACE_VALUE} = {w_sh:.1f} Cr shares",
        ],
        f"EPS = {W_PNL['pat_mi'][0]:,.1f} / {w_sh:.1f} = Rs {WR['Earnings per Share (Rs)'][0]:.2f}",
        f"EPS = Rs {WR['Earnings per Share (Rs)'][0]:.2f} per share",
        "Direct EPS comparison is limited since Wipro had a bonus issue in FY2025 doubling the share count. "
        "Over 5 years, HCL's EPS grew 56% (41.07 to 64.05) vs Wipro's per-share decline due to dilution."
    )

    # ===================== C. SOLVENCY RATIOS =====================
    _add_heading(doc, "C. Solvency Ratios", level=2)

    h_td = H_BS['lt_borr'][0] + H_BS['st_borr'][0]
    w_td = W_BS['lt_borr'][0] + W_BS['st_borr'][0]

    proof_enhanced(
        "10. Debt-Equity Ratio",
        "Total Debt / Total Shareholders' Funds\nwhere Total Debt = Long Term Borrowings + Short Term Borrowings",
        [
            f"Long Term Borrowings = {H_BS['lt_borr'][0]} (from HCL Balance Sheet \u2192 'Long Term Borrowings' row)",
            f"Short Term Borrowings = {H_BS['st_borr'][0]:,} (from HCL Balance Sheet \u2192 'Short Term Borrowings' row)",
            f"Total Debt = {H_BS['lt_borr'][0]} + {H_BS['st_borr'][0]:,} = {h_td:,}",
            f"Total Shareholders' Funds = {H_BS['total_shf'][0]:,} (from HCL Balance Sheet \u2192 'Total Shareholders Funds' row)",
        ],
        f"Debt-Equity Ratio = {h_td:,} / {H_BS['total_shf'][0]:,}",
        f"Debt-Equity Ratio = {HR['Debt-Equity Ratio'][0]:.4f}",
        [
            f"Long Term Borrowings = {W_BS['lt_borr'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Long Term Borrowings' row)",
            f"Short Term Borrowings = {W_BS['st_borr'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Short Term Borrowings' row)",
            f"Total Debt = {W_BS['lt_borr'][0]:,.1f} + {W_BS['st_borr'][0]:,.1f} = {w_td:,.1f}",
            f"Total Shareholders' Funds = {W_BS['total_shf'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Total Shareholders Funds' row)",
        ],
        f"Debt-Equity Ratio = {w_td:,.1f} / {W_BS['total_shf'][0]:,.1f}",
        f"Debt-Equity Ratio = {WR['Debt-Equity Ratio'][0]:.4f}",
        "HCL is nearly debt-free (D/E 0.03), while Wipro carries meaningful debt (D/E 0.20) "
        "primarily from acquisition financing. Lower leverage = lower financial risk."
    )

    proof_enhanced(
        "11. Interest Coverage Ratio",
        "EBIT / Finance Costs",
        [
            f"EBIT = {h_ebit:,} (derived: PBT {H_PNL['pbt'][0]:,} + Finance Costs {H_PNL['fin_cost'][0]} from HCL P&L)",
            f"Finance Costs = {H_PNL['fin_cost'][0]} (from HCL P&L \u2192 'Finance Costs' row)",
        ],
        f"Interest Coverage = {h_ebit:,} / {H_PNL['fin_cost'][0]}",
        f"Interest Coverage = {HR['Interest Coverage Ratio'][0]:.2f}x",
        [
            f"EBIT = {w_ebit:,.1f} (derived: PBT {W_PNL['pbt'][0]:,.1f} + Finance Costs {W_PNL['fin_cost'][0]:,} from Wipro P&L)",
            f"Finance Costs = {W_PNL['fin_cost'][0]:,} (from Wipro P&L \u2192 'Finance Costs' row)",
        ],
        f"Interest Coverage = {w_ebit:,.1f} / {W_PNL['fin_cost'][0]:,}",
        f"Interest Coverage = {WR['Interest Coverage Ratio'][0]:.2f}x",
        "HCL's interest coverage (37.12x) is much higher than Wipro's (12.82x). Both are comfortable, "
        "but HCL has virtually zero debt service risk."
    )

    # ===================== D. TURNOVER RATIOS =====================
    _add_heading(doc, "D. Turnover Ratios", level=2)

    h_avg_rec = (H_BS['trade_rec'][0] + H_BS['trade_rec'][1]) / 2
    w_avg_rec = (W_BS['trade_rec'][0] + W_BS['trade_rec'][1]) / 2

    proof_enhanced(
        "12. Debtors Turnover Ratio",
        "Revenue / Average Trade Receivables\nAverage = (FY2025 + FY2024) / 2",
        [
            "Revenue = 1,17,055 (from HCL P&L \u2192 'Revenue From Operations [Net]' row)",
            f"Trade Receivables FY2025 = {H_BS['trade_rec'][0]:,} (from HCL Balance Sheet FY2025 \u2192 'Trade Receivables' row)",
            f"Trade Receivables FY2024 = {H_BS['trade_rec'][1]:,} (from HCL Balance Sheet FY2024 \u2192 'Trade Receivables' row)",
            f"Average Trade Receivables = ({H_BS['trade_rec'][0]:,} + {H_BS['trade_rec'][1]:,}) / 2 = {h_avg_rec:,.1f}",
        ],
        f"Debtors Turnover = 1,17,055 / {h_avg_rec:,.1f}",
        f"Debtors Turnover = {HR['Debtors Turnover Ratio'][0]:.2f}x",
        [
            "Revenue = 89,088.4 (from Wipro P&L \u2192 'Revenue From Operations [Net]' row)",
            f"Trade Receivables FY2025 = {W_BS['trade_rec'][0]:,.1f} (from Wipro Balance Sheet FY2025 \u2192 'Trade Receivables' row)",
            f"Trade Receivables FY2024 = {W_BS['trade_rec'][1]:,.1f} (from Wipro Balance Sheet FY2024 \u2192 'Trade Receivables' row)",
            f"Average Trade Receivables = ({W_BS['trade_rec'][0]:,.1f} + {W_BS['trade_rec'][1]:,.1f}) / 2 = {w_avg_rec:,.1f}",
        ],
        f"Debtors Turnover = 89,088.4 / {w_avg_rec:,.1f}",
        f"Debtors Turnover = {WR['Debtors Turnover Ratio'][0]:.2f}x",
        "Wipro collects receivables faster (6.11x vs 4.56x), but this partly reflects "
        "Wipro's receivables normalizing from a FY2024 spike."
    )

    proof_enhanced(
        "13. Average Collection Period",
        "365 / Debtors Turnover Ratio",
        [
            f"Debtors Turnover Ratio = {HR['Debtors Turnover Ratio'][0]:.2f}x (computed above in Ratio #12)",
        ],
        f"Avg Collection Period = 365 / {HR['Debtors Turnover Ratio'][0]:.2f}",
        f"Avg Collection Period = {HR['Avg Collection Period (days)'][0]:.1f} days",
        [
            f"Debtors Turnover Ratio = {WR['Debtors Turnover Ratio'][0]:.2f}x (computed above in Ratio #12)",
        ],
        f"Avg Collection Period = 365 / {WR['Debtors Turnover Ratio'][0]:.2f}",
        f"Avg Collection Period = {WR['Avg Collection Period (days)'][0]:.1f} days",
        "HCL takes ~80 days to collect from clients vs Wipro's ~60 days. Both are within "
        "normal IT services range (60-90 days)."
    )

    h_purch = H_PNL['purchase'][0] + H_PNL['op_exp'][0] + H_PNL['other_exp'][0]
    w_purch = W_PNL['purchase'][0] + W_PNL['op_exp'][0] + W_PNL['other_exp'][0]
    h_avg_pay = (H_BS['trade_pay'][0] + H_BS['trade_pay'][1]) / 2
    w_avg_pay = (W_BS['trade_pay'][0] + W_BS['trade_pay'][1]) / 2

    proof_enhanced(
        "14. Creditors Turnover Ratio",
        "(Purchase + Operating Exp + Other Exp) / Average Trade Payables",
        [
            f"Purchase of Stock-in-Trade = {H_PNL['purchase'][0]:,} (from HCL P&L \u2192 'Purchase Of Stock-In Trade' row)",
            f"Operating & Direct Expenses = {H_PNL['op_exp'][0]:,} (from HCL P&L \u2192 'Operating And Direct Expenses' row)",
            f"Other Expenses = {H_PNL['other_exp'][0]:,} (from HCL P&L \u2192 'Other Expenses' row)",
            f"Total Purchases = {H_PNL['purchase'][0]:,} + {H_PNL['op_exp'][0]:,} + {H_PNL['other_exp'][0]:,} = {h_purch:,}",
            f"Trade Payables FY2025 = {H_BS['trade_pay'][0]:,} (from HCL Balance Sheet \u2192 'Trade Payables' row)",
            f"Trade Payables FY2024 = {H_BS['trade_pay'][1]:,} (from HCL Balance Sheet FY2024 \u2192 same row)",
            f"Average Trade Payables = ({H_BS['trade_pay'][0]:,} + {H_BS['trade_pay'][1]:,}) / 2 = {h_avg_pay:,.1f}",
        ],
        f"Creditors Turnover = {h_purch:,} / {h_avg_pay:,.1f}",
        f"Creditors Turnover = {HR['Creditors Turnover Ratio'][0]:.2f}x",
        [
            f"Purchase of Stock-in-Trade = {W_PNL['purchase'][0]} (from Wipro P&L \u2192 'Purchase Of Stock-In Trade' row)",
            f"Operating & Direct Expenses = {W_PNL['op_exp'][0]:,.1f} (from Wipro P&L \u2192 'Operating And Direct Expenses' row)",
            f"Other Expenses = {W_PNL['other_exp'][0]:,.1f} (from Wipro P&L \u2192 'Other Expenses' row)",
            f"Total Purchases = {W_PNL['purchase'][0]} + {W_PNL['op_exp'][0]:,.1f} + {W_PNL['other_exp'][0]:,.1f} = {w_purch:,.1f}",
            f"Trade Payables FY2025 = {W_BS['trade_pay'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Trade Payables' row)",
            f"Trade Payables FY2024 = {W_BS['trade_pay'][1]:,.1f} (from Wipro Balance Sheet FY2024 \u2192 same row)",
            f"Average Trade Payables = ({W_BS['trade_pay'][0]:,.1f} + {W_BS['trade_pay'][1]:,.1f}) / 2 = {w_avg_pay:,.1f}",
        ],
        f"Creditors Turnover = {w_purch:,.1f} / {w_avg_pay:,.1f}",
        f"Creditors Turnover = {WR['Creditors Turnover Ratio'][0]:.2f}x",
        "Lower creditors turnover for HCL (2.59x vs 3.02x) means HCL takes longer to pay suppliers, "
        "which is favorable for cash management but may reflect the FY2025 payables spike."
    )

    proof_enhanced(
        "15. Average Payment Period",
        "365 / Creditors Turnover Ratio",
        [
            f"Creditors Turnover Ratio = {HR['Creditors Turnover Ratio'][0]:.2f}x (computed above in Ratio #14)",
        ],
        f"Avg Payment Period = 365 / {HR['Creditors Turnover Ratio'][0]:.2f}",
        f"Avg Payment Period = {HR['Avg Payment Period (days)'][0]:.1f} days",
        [
            f"Creditors Turnover Ratio = {WR['Creditors Turnover Ratio'][0]:.2f}x (computed above in Ratio #14)",
        ],
        f"Avg Payment Period = 365 / {WR['Creditors Turnover Ratio'][0]:.2f}",
        f"Avg Payment Period = {WR['Avg Payment Period (days)'][0]:.1f} days",
        "HCL takes ~141 days vs Wipro's ~121 days to pay creditors."
    )

    h_cogs_v = H_PNL['purchase'][0] + H_PNL['op_exp'][0] + H_PNL['inv_change'][0]
    w_cogs_v = W_PNL['purchase'][0] + W_PNL['op_exp'][0] + W_PNL['inv_change'][0]
    h_avg_inv = (H_BS['inventories'][0] + H_BS['inventories'][1]) / 2
    w_avg_inv = (W_BS['inventories'][0] + W_BS['inventories'][1]) / 2

    proof_enhanced(
        "16. Inventory Turnover Ratio",
        "COGS / Average Inventories\nwhere COGS = Purchase + Operating Exp + Inventory Changes",
        [
            f"COGS = {h_cogs_v:,} (derived: {H_PNL['purchase'][0]:,} + {H_PNL['op_exp'][0]:,} + {H_PNL['inv_change'][0]} from HCL P&L)",
            f"Inventories FY2025 = {H_BS['inventories'][0]} (from HCL Balance Sheet \u2192 'Inventories' row)",
            f"Inventories FY2024 = {H_BS['inventories'][1]} (from HCL Balance Sheet FY2024 \u2192 same row)",
            f"Average Inventories = ({H_BS['inventories'][0]} + {H_BS['inventories'][1]}) / 2 = {h_avg_inv:.1f}",
        ],
        f"Inventory Turnover = {h_cogs_v:,} / {h_avg_inv:.1f}",
        f"Inventory Turnover = {HR['Inventory Turnover Ratio'][0]:.2f}x",
        [
            f"COGS = {w_cogs_v:,.1f} (derived: {W_PNL['purchase'][0]} + {W_PNL['op_exp'][0]:,.1f} + {W_PNL['inv_change'][0]} from Wipro P&L)",
            f"Inventories FY2025 = {W_BS['inventories'][0]} (from Wipro Balance Sheet \u2192 'Inventories' row)",
            f"Inventories FY2024 = {W_BS['inventories'][1]} (from Wipro Balance Sheet FY2024 \u2192 same row)",
            f"Average Inventories = ({W_BS['inventories'][0]} + {W_BS['inventories'][1]}) / 2 = {w_avg_inv:.1f}",
        ],
        f"Inventory Turnover = {w_cogs_v:,.1f} / {w_avg_inv:.1f}",
        f"Inventory Turnover = {WR['Inventory Turnover Ratio'][0]:.2f}x",
        "Extremely high turnover for both companies (100x+) because IT services firms hold minimal "
        "inventory. This ratio is less meaningful for the services sector."
    )

    proof_enhanced(
        "17. Inventory Holding Days",
        "365 / Inventory Turnover Ratio",
        [
            f"Inventory Turnover Ratio = {HR['Inventory Turnover Ratio'][0]:.2f}x (computed above in Ratio #16)",
        ],
        f"Inventory Holding Days = 365 / {HR['Inventory Turnover Ratio'][0]:.2f}",
        f"Inventory Holding Days = {HR['Inventory Holding Days'][0]:.1f} days",
        [
            f"Inventory Turnover Ratio = {WR['Inventory Turnover Ratio'][0]:.2f}x (computed above in Ratio #16)",
        ],
        f"Inventory Holding Days = 365 / {WR['Inventory Turnover Ratio'][0]:.2f}",
        f"Inventory Holding Days = {WR['Inventory Holding Days'][0]:.1f} days",
        "Near-zero holding days confirm the asset-light IT services model."
    )

    h_avg_fa = (H_BS['fixed_assets'][0] + H_BS['fixed_assets'][1]) / 2
    w_avg_fa = (W_BS['fixed_assets'][0] + W_BS['fixed_assets'][1]) / 2

    proof_enhanced(
        "18. Fixed Asset Turnover Ratio",
        "Revenue / Average Fixed Assets\nAverage = (FY2025 + FY2024) / 2",
        [
            "Revenue = 1,17,055 (from HCL P&L \u2192 'Revenue From Operations [Net]' row)",
            f"Fixed Assets FY2025 = {H_BS['fixed_assets'][0]:,} (from HCL Balance Sheet \u2192 'Fixed Assets' row)",
            f"Fixed Assets FY2024 = {H_BS['fixed_assets'][1]:,} (from HCL Balance Sheet FY2024 \u2192 same row)",
            f"Average Fixed Assets = ({H_BS['fixed_assets'][0]:,} + {H_BS['fixed_assets'][1]:,}) / 2 = {h_avg_fa:,.1f}",
        ],
        f"Fixed Asset Turnover = 1,17,055 / {h_avg_fa:,.1f}",
        f"Fixed Asset Turnover = {HR['Fixed Asset Turnover'][0]:.2f}x",
        [
            "Revenue = 89,088.4 (from Wipro P&L \u2192 'Revenue From Operations [Net]' row)",
            f"Fixed Assets FY2025 = {W_BS['fixed_assets'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Fixed Assets' row)",
            f"Fixed Assets FY2024 = {W_BS['fixed_assets'][1]:,.1f} (from Wipro Balance Sheet FY2024 \u2192 same row)",
            f"Average Fixed Assets = ({W_BS['fixed_assets'][0]:,.1f} + {W_BS['fixed_assets'][1]:,.1f}) / 2 = {w_avg_fa:,.1f}",
        ],
        f"Fixed Asset Turnover = 89,088.4 / {w_avg_fa:,.1f}",
        f"Fixed Asset Turnover = {WR['Fixed Asset Turnover'][0]:.2f}x",
        "HCL generates Rs 7.93 of revenue per Rs 1 of fixed assets vs Wipro's Rs 6.71. "
        "Higher asset turnover reflects HCL's more efficient use of physical infrastructure."
    )

    # ===================== E. VALUATION RATIO =====================
    _add_heading(doc, "E. Valuation Ratio", level=2)

    proof_enhanced(
        "19. P/E Ratio",
        "Market Price per Share / EPS\nMarket Price = Price/BV x Book Value per Share (from Moneycontrol Ratios page)",
        [
            f"Price/BV = {H_MKT['price_bv'][0]} (from HCL Moneycontrol \u2192 'Price/BV (X)' ratio)",
            f"Book Value per Share = {H_MKT['bv_share'][0]} (from HCL Moneycontrol \u2192 'Book Value/Share' ratio)",
            f"Market Price per Share = {H_MKT['price_bv'][0]} x {H_MKT['bv_share'][0]} = {HR['_mp'][0]:.2f}",
            f"EPS = {HR['Earnings per Share (Rs)'][0]:.2f} (computed above in Ratio #9)",
        ],
        f"P/E Ratio = {HR['_mp'][0]:.2f} / {HR['Earnings per Share (Rs)'][0]:.2f}",
        f"P/E Ratio = {HR['P/E Ratio'][0]:.2f}x",
        [
            f"Price/BV = {W_MKT['price_bv'][0]} (from Wipro Moneycontrol \u2192 'Price/BV (X)' ratio)",
            f"Book Value per Share = {W_MKT['bv_share'][0]} (from Wipro Moneycontrol \u2192 'Book Value/Share' ratio)",
            f"Market Price per Share = {W_MKT['price_bv'][0]} x {W_MKT['bv_share'][0]} = {WR['_mp'][0]:.2f}",
            f"EPS = {WR['Earnings per Share (Rs)'][0]:.2f} (computed above in Ratio #9)",
        ],
        f"P/E Ratio = {WR['_mp'][0]:.2f} / {WR['Earnings per Share (Rs)'][0]:.2f}",
        f"P/E Ratio = {WR['P/E Ratio'][0]:.2f}x",
        "HCL commands a higher P/E ({:.1f}x vs {:.1f}x), reflecting the market's willingness to pay "
        "a premium for HCL's consistent growth and superior return profile.".format(
            HR['P/E Ratio'][0], WR['P/E Ratio'][0])
    )

    # ===================== F. DUPONT ANALYSIS PROOF =====================
    _add_heading(doc, "F. DuPont Analysis - Detailed Proof (FY2025)", level=2)

    _add_para(doc, "Formula: ROE = Net Profit Margin x Asset Turnover x Equity Multiplier x 100", bold=True, size=11)
    _add_para(doc, "")

    _add_para(doc, "HCL Technologies FY2025:", bold=True, size=11)
    _add_para(doc, "Step 1: Net Profit Margin (NPM) = PAT / Revenue", bold=True, size=10)
    for inp in [
        f"PAT = {H_PNL['pat'][0]:,} (from HCL P&L \u2192 'Profit/Loss After Tax' row)",
        f"Revenue = {H_PNL['revenue'][0]:,} (from HCL P&L \u2192 'Revenue From Operations [Net]' row)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"NPM = {H_PNL['pat'][0]:,} / {H_PNL['revenue'][0]:,} = {HR['_npm'][0]:.4f}", size=10)

    _add_para(doc, "Step 2: Asset Turnover (AT) = Revenue / Total Assets", bold=True, size=10)
    for inp in [
        f"Revenue = {H_PNL['revenue'][0]:,} (from HCL P&L \u2192 'Revenue From Operations [Net]' row)",
        f"Total Assets = {H_BS['total_assets'][0]:,} (from HCL Balance Sheet \u2192 'Total Assets' row)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"AT = {H_PNL['revenue'][0]:,} / {H_BS['total_assets'][0]:,} = {HR['_at'][0]:.4f}", size=10)

    _add_para(doc, "Step 3: Equity Multiplier (EM) = Total Assets / Total Shareholders' Funds", bold=True, size=10)
    for inp in [
        f"Total Assets = {H_BS['total_assets'][0]:,} (from HCL Balance Sheet \u2192 'Total Assets' row)",
        f"Total Shareholders' Funds = {H_BS['total_shf'][0]:,} (from HCL Balance Sheet \u2192 'Total Shareholders Funds' row)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"EM = {H_BS['total_assets'][0]:,} / {H_BS['total_shf'][0]:,} = {HR['_em'][0]:.4f}", size=10)

    _add_para(doc, f"Step 4: ROE = NPM x AT x EM x 100 = {HR['_npm'][0]:.4f} x {HR['_at'][0]:.4f} x {HR['_em'][0]:.4f} x 100 = {HR['_roe_d'][0]:.2f}%", bold=True, size=10)
    _add_para(doc, "")

    _add_para(doc, "Wipro FY2025:", bold=True, size=11)
    _add_para(doc, "Step 1: Net Profit Margin (NPM) = PAT / Revenue", bold=True, size=10)
    for inp in [
        f"PAT = {W_PNL['pat'][0]:,.1f} (from Wipro P&L \u2192 'Profit/Loss After Tax' row)",
        f"Revenue = {W_PNL['revenue'][0]:,.1f} (from Wipro P&L \u2192 'Revenue From Operations [Net]' row)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"NPM = {W_PNL['pat'][0]:,.1f} / {W_PNL['revenue'][0]:,.1f} = {WR['_npm'][0]:.4f}", size=10)

    _add_para(doc, "Step 2: Asset Turnover (AT) = Revenue / Total Assets", bold=True, size=10)
    for inp in [
        f"Revenue = {W_PNL['revenue'][0]:,.1f} (from Wipro P&L \u2192 'Revenue From Operations [Net]' row)",
        f"Total Assets = {W_BS['total_assets'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Total Assets' row)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"AT = {W_PNL['revenue'][0]:,.1f} / {W_BS['total_assets'][0]:,.1f} = {WR['_at'][0]:.4f}", size=10)

    _add_para(doc, "Step 3: Equity Multiplier (EM) = Total Assets / Total Shareholders' Funds", bold=True, size=10)
    for inp in [
        f"Total Assets = {W_BS['total_assets'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Total Assets' row)",
        f"Total Shareholders' Funds = {W_BS['total_shf'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Total Shareholders Funds' row)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"EM = {W_BS['total_assets'][0]:,.1f} / {W_BS['total_shf'][0]:,.1f} = {WR['_em'][0]:.4f}", size=10)

    _add_para(doc, f"Step 4: ROE = NPM x AT x EM x 100 = {WR['_npm'][0]:.4f} x {WR['_at'][0]:.4f} x {WR['_em'][0]:.4f} x 100 = {WR['_roe_d'][0]:.2f}%", bold=True, size=10)
    _add_para(doc, "")

    # ===================== G. CCC ANALYSIS PROOF =====================
    _add_heading(doc, "G. Cash Conversion Cycle - Detailed Proof (FY2025)", level=2)

    _add_para(doc, "Formula: CCC = DIO + DSO - DPO (all in days)", bold=True, size=11)
    _add_para(doc, "")

    _add_para(doc, "HCL Technologies FY2025:", bold=True, size=11)

    _add_para(doc, "Step 1: Days Inventory Outstanding (DIO) = 365 x Average Inventories / COGS", bold=True, size=10)
    for inp in [
        f"Inventories FY2025 = {H_BS['inventories'][0]} (from HCL Balance Sheet \u2192 'Inventories' row)",
        f"Inventories FY2024 = {H_BS['inventories'][1]} (from HCL Balance Sheet FY2024 \u2192 same row)",
        f"Average Inventories = ({H_BS['inventories'][0]} + {H_BS['inventories'][1]}) / 2 = {h_avg_inv:.1f}",
        f"COGS = {h_cogs_v:,} (Purchase {H_PNL['purchase'][0]:,} + Op Exp {H_PNL['op_exp'][0]:,} + Inv Change {H_PNL['inv_change'][0]} from HCL P&L)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"DIO = 365 x {h_avg_inv:.1f} / {h_cogs_v:,} = {HR['_dio'][0]:.1f} days", size=10)

    h_avg_rec2 = (H_BS['trade_rec'][0] + H_BS['trade_rec'][1]) / 2
    _add_para(doc, "Step 2: Days Sales Outstanding (DSO) = 365 x Average Trade Receivables / Revenue", bold=True, size=10)
    for inp in [
        f"Trade Receivables FY2025 = {H_BS['trade_rec'][0]:,} (from HCL Balance Sheet \u2192 'Trade Receivables' row)",
        f"Trade Receivables FY2024 = {H_BS['trade_rec'][1]:,} (from HCL Balance Sheet FY2024 \u2192 same row)",
        f"Average Trade Receivables = ({H_BS['trade_rec'][0]:,} + {H_BS['trade_rec'][1]:,}) / 2 = {h_avg_rec2:,.1f}",
        f"Revenue = {H_PNL['revenue'][0]:,} (from HCL P&L \u2192 'Revenue From Operations [Net]' row)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"DSO = 365 x {h_avg_rec2:,.1f} / {H_PNL['revenue'][0]:,} = {HR['_dso'][0]:.1f} days", size=10)

    h_avg_pay2 = (H_BS['trade_pay'][0] + H_BS['trade_pay'][1]) / 2
    _add_para(doc, "Step 3: Days Payable Outstanding (DPO) = 365 x Average Trade Payables / Purchases", bold=True, size=10)
    for inp in [
        f"Trade Payables FY2025 = {H_BS['trade_pay'][0]:,} (from HCL Balance Sheet \u2192 'Trade Payables' row)",
        f"Trade Payables FY2024 = {H_BS['trade_pay'][1]:,} (from HCL Balance Sheet FY2024 \u2192 same row)",
        f"Average Trade Payables = ({H_BS['trade_pay'][0]:,} + {H_BS['trade_pay'][1]:,}) / 2 = {h_avg_pay2:,.1f}",
        f"Purchases = {h_purch:,} (Purchase {H_PNL['purchase'][0]:,} + Op Exp {H_PNL['op_exp'][0]:,} + Other Exp {H_PNL['other_exp'][0]:,} from HCL P&L)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"DPO = 365 x {h_avg_pay2:,.1f} / {h_purch:,} = {HR['_dpo'][0]:.1f} days", size=10)

    _add_para(doc, f"Step 4: CCC = DIO + DSO - DPO = {HR['_dio'][0]:.1f} + {HR['_dso'][0]:.1f} - {HR['_dpo'][0]:.1f} = {HR['_ccc'][0]:.1f} days", bold=True, size=10)
    _add_para(doc, "")

    _add_para(doc, "Wipro FY2025:", bold=True, size=11)

    _add_para(doc, "Step 1: Days Inventory Outstanding (DIO) = 365 x Average Inventories / COGS", bold=True, size=10)
    for inp in [
        f"Inventories FY2025 = {W_BS['inventories'][0]} (from Wipro Balance Sheet \u2192 'Inventories' row)",
        f"Inventories FY2024 = {W_BS['inventories'][1]} (from Wipro Balance Sheet FY2024 \u2192 same row)",
        f"Average Inventories = ({W_BS['inventories'][0]} + {W_BS['inventories'][1]}) / 2 = {w_avg_inv:.1f}",
        f"COGS = {w_cogs_v:,.1f} (Purchase {W_PNL['purchase'][0]} + Op Exp {W_PNL['op_exp'][0]:,.1f} + Inv Change {W_PNL['inv_change'][0]} from Wipro P&L)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"DIO = 365 x {w_avg_inv:.1f} / {w_cogs_v:,.1f} = {WR['_dio'][0]:.1f} days", size=10)

    w_avg_rec2 = (W_BS['trade_rec'][0] + W_BS['trade_rec'][1]) / 2
    _add_para(doc, "Step 2: Days Sales Outstanding (DSO) = 365 x Average Trade Receivables / Revenue", bold=True, size=10)
    for inp in [
        f"Trade Receivables FY2025 = {W_BS['trade_rec'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Trade Receivables' row)",
        f"Trade Receivables FY2024 = {W_BS['trade_rec'][1]:,.1f} (from Wipro Balance Sheet FY2024 \u2192 same row)",
        f"Average Trade Receivables = ({W_BS['trade_rec'][0]:,.1f} + {W_BS['trade_rec'][1]:,.1f}) / 2 = {w_avg_rec2:,.1f}",
        f"Revenue = {W_PNL['revenue'][0]:,.1f} (from Wipro P&L \u2192 'Revenue From Operations [Net]' row)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"DSO = 365 x {w_avg_rec2:,.1f} / {W_PNL['revenue'][0]:,.1f} = {WR['_dso'][0]:.1f} days", size=10)

    w_avg_pay2 = (W_BS['trade_pay'][0] + W_BS['trade_pay'][1]) / 2
    _add_para(doc, "Step 3: Days Payable Outstanding (DPO) = 365 x Average Trade Payables / Purchases", bold=True, size=10)
    for inp in [
        f"Trade Payables FY2025 = {W_BS['trade_pay'][0]:,.1f} (from Wipro Balance Sheet \u2192 'Trade Payables' row)",
        f"Trade Payables FY2024 = {W_BS['trade_pay'][1]:,.1f} (from Wipro Balance Sheet FY2024 \u2192 same row)",
        f"Average Trade Payables = ({W_BS['trade_pay'][0]:,.1f} + {W_BS['trade_pay'][1]:,.1f}) / 2 = {w_avg_pay2:,.1f}",
        f"Purchases = {w_purch:,.1f} (Purchase {W_PNL['purchase'][0]} + Op Exp {W_PNL['op_exp'][0]:,.1f} + Other Exp {W_PNL['other_exp'][0]:,.1f} from Wipro P&L)",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(inp)
        run.font.name = 'Times New Roman'; run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    _add_para(doc, f"DPO = 365 x {w_avg_pay2:,.1f} / {w_purch:,.1f} = {WR['_dpo'][0]:.1f} days", size=10)

    _add_para(doc, f"Step 4: CCC = DIO + DSO - DPO = {WR['_dio'][0]:.1f} + {WR['_dso'][0]:.1f} - {WR['_dpo'][0]:.1f} = {WR['_ccc'][0]:.1f} days", bold=True, size=10)

    doc.save(out_path)
    print(f"[OK] Ratio proofs saved: {out_path}")


# ==========================  MAIN  ==========================
def main():
    FINAL.mkdir(parents=True, exist_ok=True)

    print("=== Step 1: Cleaning dataset files (removing standalone, removing Ratios sheet) ===")
    clean_dataset_files()

    print("\n=== Step 2: Building Excel report (full data + ratios) ===")
    excel_path = FINAL / "excel_report.xlsx"
    create_excel(excel_path)

    print("\n=== Step 3: Building Word report ===")
    report_path = FINAL / "word_report.docx"
    create_word_report(report_path)

    print("\n=== Step 4: Building Ratio Calculation Proofs (enhanced with source derivations) ===")
    proofs_path = FINAL / "ratio_calculation_proofs.docx"
    create_ratio_proofs(proofs_path)

    print("\n=== ALL DONE ===")
    print(f"Output folder: {FINAL}")
    for f in sorted(FINAL.iterdir()):
        if f.suffix in ['.xlsx', '.docx', '.py', '.md']:
            print(f"  {f.name}")

if __name__ == "__main__":
    main()
